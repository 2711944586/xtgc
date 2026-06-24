from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

from .config import DATA_DIR, FIGURES_DIR, REPORTS_DIR, SCREENSHOTS_DIR, TABLES_DIR, ensure_dirs
from .utils import read_json


NAVY = RGBColor(22, 58, 70)
TEAL = RGBColor(20, 124, 124)
ORANGE = RGBColor(201, 111, 45)
INK = RGBColor(23, 35, 38)
MUTED = RGBColor(88, 104, 109)
LINE = "CFDAD5"
SOFT = "EDF2EF"
CODE_BG = "11272E"


def _read_table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES_DIR / name)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(8.5)
    if color is not None:
        run.font.color.rgb = color


def _set_paragraph_border(paragraph, color: str = "D7E1DE") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 15, NAVY, 14, 7),
        ("Heading 2", 12.5, TEAL, 10, 5),
        ("Heading 3", 11.2, ORANGE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "FlightResilience | 航空网络延误传播与恢复策略"
    header.runs[0].font.name = "Microsoft YaHei"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "数据来源：U.S. DOT BTS TranStats；仿真成本为相对情景参数"
    footer.runs[0].font.name = "Microsoft YaHei"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def _add_title(doc: Document, audit: dict, decision: dict) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("SYSTEM ENGINEERING COURSE PROJECT")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(9)
    r.font.color.rgb = TEAL

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("面向突发扰动的航空网络延误传播识别与恢复策略决策")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(19)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("基于 BTS 公开数据的系统工程建模、状态传播仿真与静态 Web 演示")
    sr.font.name = "Microsoft YaHei"
    sr.font.size = Pt(11)
    sr.font.color.rgb = MUTED

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    facts = [
        ("样本范围", f"{audit['date_min']} 至 {audit['date_max']}"),
        ("航班记录", f"{audit['scoped_rows_top_airports']:,}"),
        ("机场节点", "15"),
        ("主推荐", decision["recommended_strategy"]),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, facts):
        _set_cell_shading(cell, SOFT)
        _set_cell_text(cell, f"{label}\n{value}", bold=True, color=NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("课程：系统工程    项目：FlightResilience    输出：报告 / PPT / 讲稿 / 静态 Web Demo").font.size = Pt(9)
    _set_paragraph_border(p)


def _add_body(doc: Document, text: str) -> None:
    for para in text.strip().split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())


def _add_formula(doc: Document, formula: str, note: str) -> None:
    p = doc.add_paragraph()
    _shade_paragraph(p, "F7E7D9")
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(formula)
    r.bold = True
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = NAVY
    n = p.add_run(f"\n{note}")
    n.font.name = "Microsoft YaHei"
    n.font.size = Pt(9)
    n.font.color.rgb = MUTED


def _add_code_block(doc: Document, title: str, code: str) -> None:
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(5)
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.color.rgb = TEAL
    p = doc.add_paragraph()
    _shade_paragraph(p, CODE_BG)
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(237, 247, 242)


def _add_reading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _shade_paragraph(p, SOFT)
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("读图解释：")
    r.bold = True
    r.font.color.rgb = TEAL
    p.add_run(text)


def _add_df_table(doc: Document, df: pd.DataFrame, caption: str, max_rows: int = 8) -> None:
    view = df.head(max_rows).copy()
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.bold = True
    cap.runs[0].font.color.rgb = NAVY
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(view.columns):
        _set_cell_shading(hdr[i], "163A46")
        _set_cell_text(hdr[i], str(col), bold=True, color=RGBColor(255, 255, 255))
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if isinstance(val, float):
                text = f"{val:.4g}"
            else:
                text = str(val)
            _set_cell_text(cells[i], text)
    doc.add_paragraph()


def _add_picture(doc: Document, name: str, caption: str, reading: str, width: float = 5.95, source_dir: Path = FIGURES_DIR) -> None:
    path = source_dir / name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.name = "Microsoft YaHei"
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = MUTED
    _add_reading(doc, reading)


def _format_strategy_name(value: str) -> str:
    return {
        "baseline": "基准策略",
        "uniform_buffer": "统一缓冲",
        "hub_priority": "关键枢纽优先",
        "dynamic_combo": "动态组合",
    }.get(value, value)


def generate_docx_report() -> Path:
    ensure_dirs()
    audit = read_json(DATA_DIR / "data_audit.json")
    model = read_json(REPORTS_DIR.parent / "models" / "model_summary.json")
    prop = read_json(TABLES_DIR / "propagation_validation.json")
    decision = read_json(REPORTS_DIR.parent / "data" / "demo" / "decision_summary.json")
    nodes = _read_table("airport_nodes.csv")
    rankings = _read_table("strategy_rankings.csv")
    weights = _read_table("indicator_weights.csv")
    risk = _read_table("risk_decision.csv")
    metrics_by_scenario = _read_table("strategy_metrics_by_scenario.csv")
    manifest = pd.read_csv(DATA_DIR / "data_manifest.csv")

    doc = Document()
    _style_doc(doc)
    _add_title(doc, audit, decision)

    doc.add_heading("摘要", level=1)
    _add_body(
        doc,
        f"""
        本项目面向航空网络突发扰动下的延误传播与恢复策略选择问题，使用美国交通部 BTS TranStats 航班准点数据，构建“真实运行数据 - 计划阶段延误预测 - 机场复杂网络 - ISM 结构解释 - 状态空间传播仿真 - 多准则恢复决策 - 静态 Web 演示”的系统工程闭环。样本覆盖 {audit['date_min']} 至 {audit['date_max']}，筛选样本内总流量前 15 个机场，共 {audit['scoped_rows_top_airports']:,} 条记录。

        研究结论表明，延误不是均匀发生的随机噪声，而是具有显著时间异质性、机场差异和网络位置效应。计划阶段预测模型在时间外测试集上的 ROC-AUC 为 {model['best_test_metrics']['roc_auc']:.3f}，其主要价值不是单独替代调度，而是为传播矩阵和恢复资源优先级提供风险输入。传播矩阵估计后谱半径为 {prop['spectral_radius_final']:.3f}，说明状态递推在本项目情景内保持稳定。

        在主偏好权重下，{_format_strategy_name(decision['recommended_strategy'])} 的 TOPSIS 综合排名第一；但风险型决策和权重敏感性分析显示，若极端强调低成本或保守损失，{_format_strategy_name(decision['risk_recommended_strategy'])} 会反超。因此本文最终给出“有条件推荐”，而不是把某一策略表述为无条件最优。
        """,
    )
    doc.add_paragraph("关键词：航空延误；复杂网络；系统工程；状态空间；TOPSIS；风险决策；静态 Web Demo")

    doc.add_heading("1 问题背景与研究对象", level=1)
    _add_body(
        doc,
        """
        航班延误常被理解为单个航班的准点性问题，但在枢纽化运行和航班串飞背景下，一个机场的容量下降、前序航班晚到或局部天气冲击，可能通过航线连接和机组/飞机轮转在数小时内扩散到其他机场。若只预测“某个航班是否会延误”，仍无法回答恢复决策最关心的三个问题：延误会传到哪里，哪些机场应优先恢复，有限资源如何分配。

        系统工程视角的价值在于将问题从“分类预测”提升为“网络恢复决策”。本文把机场、航线、航班计划、航空公司、运行状态和恢复资源共同视为航空运行系统；外部扰动包括天气、高峰需求、空域限制和枢纽容量下降。系统目标不是单一最小延误，而是在运行效率、网络韧性、航班影响、恢复成本和可实施性之间形成可解释权衡。
        """,
    )
    _add_formula(doc, "p_i = P(y_i = 1 | x_i) = f(x_i)", "p_i 为计划阶段到达延误概率，y_i=1 表示 ArrDel15；x_i 只包含起飞前可获得或可由训练期统计得到的特征。")

    doc.add_heading("2 系统工程方法链", level=1)
    _add_body(
        doc,
        """
        本项目以霍尔三维结构组织技术路线：逻辑维完成问题定义、模型建立、优化与决策；时间维覆盖数据获取、模型训练、仿真评估和交付部署；知识维整合统计学习、复杂网络、系统动力学和综合评价。切克兰德软系统方法用于补充解释不同主体目标冲突，例如乘客希望少延误，航空公司关注成本，机场关注容量恢复，监管者关注安全和公平。

        具体方法链包括：用鱼骨图识别延误影响因素；用 ISM 描述因素间结构关系；用时间切分模型估计计划阶段风险；用有向加权网络识别关键机场；用带网络掩码的状态空间模型模拟延误传播；用 AHP、熵权、TOPSIS、模糊综合评价、风险型和不确定型决策输出策略建议。
        """,
    )
    _add_formula(doc, "K_i = omega_1 flow_i + omega_2 betweenness_i + omega_3 risk_i + omega_4 delay_i", "K_i 为机场综合关键性，用于把流量、网络位置、预测风险和历史延误合成同一决策指标。")
    _add_picture(doc, "fig_25_ism_hierarchy.png", "图 1  延误因素多级递阶结构", "ISM 将天气/容量/计划等根源因素与拥堵、前序晚到、传播范围和恢复时间连接起来，使后续模型解释不只停留在特征重要性层面。")

    doc.add_heading("3 数据来源、审计与防泄漏处理", level=1)
    _add_body(
        doc,
        f"""
        数据来自 U.S. DOT BTS Airline On-Time Performance 月度公开文件。本文选取 2024 年 1 月至 3 月，按样本内总流量筛选前 15 个机场。清洗时保留航班计划、机场、航司、距离、计划飞行时间、取消标记和到达延误结果，同时构造日期、星期、小时、航线历史延误率、机场历史延误率、滚动拥堵代理等变量。

        防泄漏是预测可信度的核心。本项目采用时间顺序划分，而不是随机切分；历史统计特征仅由训练窗口估计，再映射到验证和测试窗口；实际起飞、实际到达、到达延误分钟数、延误原因和事后滑行信息不得进入计划阶段主模型。这样牺牲一部分表面指标，但使模型更接近真实决策场景。
        """,
    )
    _add_df_table(doc, manifest[["file_name", "rows", "columns", "sha256"]], "表 1  数据资产清单节选", max_rows=5)
    _add_code_block(
        doc,
        "代码片段 1  时间切分与历史特征映射",
        """
train = df[df["FlightDate"] < valid_start].copy()
valid = df[(df["FlightDate"] >= valid_start) & (df["FlightDate"] < test_start)].copy()
test = df[df["FlightDate"] >= test_start].copy()

route_rate = train.groupby("route")["ArrDel15"].mean()
airport_rate = train.groupby("Origin")["ArrDel15"].mean()
for part in [train, valid, test]:
    part["route_hist_delay_rate"] = part["route"].map(route_rate).fillna(train["ArrDel15"].mean())
    part["origin_hist_delay_rate"] = part["Origin"].map(airport_rate).fillna(train["ArrDel15"].mean())
        """,
    )
    _add_picture(doc, "fig_13_split_delay_rate.png", "图 2  训练/验证/测试目标比例", "时间外测试集的目标比例与训练期不同，说明随机切分会高估稳定性；后续结果均以时间外表现为主要依据。")

    doc.add_heading("4 探索性分析：延误并非均匀发生", level=1)
    _add_body(
        doc,
        """
        描述性分析显示，延误分钟数呈长尾分布，少量严重延误会显著拉高平均值；延误率在星期、小时和机场之间具有明显异质性。尤其在高峰时段和部分枢纽机场，延误风险呈现聚集特征。这个现象意味着恢复策略不能只按平均水平配置资源，而应识别何时、何地、通过哪条网络路径发生放大。
        """,
    )
    _add_picture(doc, "fig_06_delay_distribution.png", "图 3  到达延误分钟数分布", "主体分布集中在较低延误区间，但右尾很长。恢复策略的价值主要体现在压缩严重延误尾部，而不是只改善平均值。")
    _add_picture(doc, "fig_07_daily_volume_delay.png", "图 4  每日航班量与延误率趋势", "航班量与延误率并不总是同步变化，说明需求强度之外还存在天气、容量、计划缓冲和网络传导因素。")
    _add_picture(doc, "fig_08_week_hour_heatmap.png", "图 5  星期与小时延误率热力图", "热力图用于定位高风险时段；它也解释了为什么模型需要计划小时和历史拥堵类特征。")
    _add_picture(doc, "fig_09_volume_delay_scatter.png", "图 6  航班量与延误率关系", "高流量机场不必然具有最高延误率，恢复优先级必须综合流量、风险和网络中心性。")
    _add_picture(doc, "fig_10_airport_delay_rank.png", "图 7  主要机场延误率排名", "机场间差异构成后续节点画像的基础，也为 Web Demo 的机场选择联动提供解释。")

    doc.add_heading("5 计划阶段延误预测模型", level=1)
    _add_body(
        doc,
        f"""
        预测目标为 ArrDel15，即到达延误是否达到 15 分钟。本文比较逻辑回归、随机森林和 LightGBM，并以时间外 ROC-AUC、PR-AUC、F1、召回率和 Brier 分数综合判断。最佳模型为 {model['best_model']}，测试 ROC-AUC 为 {model['best_test_metrics']['roc_auc']:.3f}，PR-AUC 为 {model['best_test_metrics']['pr_auc']:.3f}。由于延误样本比例较低，PR-AUC 和召回率比单纯准确率更能反映实际使用价值。

        模型解释采用 SHAP 全局重要性。结果显示航线历史延误率、机场小时历史延误率、月份/时段和目的地相关特征贡献较高。这里需要强调：预测概率不是最终结论，而是状态传播模型和动态组合策略的输入。
        """,
    )
    _add_formula(doc, "Brier = (1/n) * sum_i (p_i - y_i)^2", "Brier 分数用于检查概率输出校准程度；调度决策更需要可比较概率，而不仅是 0/1 标签。")
    _add_picture(doc, "fig_14_model_metrics.png", "图 8  模型指标对比", "随机森林在 ROC-AUC、PR-AUC 和 F1 上取得较均衡表现，因此作为后续风险输入的主模型。")
    _add_picture(doc, "fig_15_confusion_matrix.png", "图 9  最佳模型混淆矩阵", "混淆矩阵展示阈值下的误报和漏报；恢复策略中更关注高风险样本排序和下游传播影响。")
    _add_picture(doc, "fig_17_calibration_curve.png", "图 10  概率校准曲线", "若概率严重失真，后续策略优先级会被放大误差；校准曲线用于检查风险输入的可解释性。")
    _add_picture(doc, "fig_18_shap_summary.png", "图 11  SHAP 特征重要性", "历史运行状态和航线经验风险贡献最高，说明计划阶段仍能利用过去信息形成有价值的风险估计。")
    _add_code_block(
        doc,
        "代码片段 2  模型训练与时间外评估",
        """
model.fit(X_train, y_train)
prob_valid = model.predict_proba(X_valid)[:, 1]
threshold = choose_threshold(prob_valid, y_valid, target="f1")
prob_test = model.predict_proba(X_test)[:, 1]
metrics = evaluate_binary(y_test, prob_test, threshold=threshold)
        """,
    )

    doc.add_heading("6 机场网络、关键节点与传播位置", level=1)
    top_airports = "、".join(nodes.sort_values("criticality", ascending=False).head(5)["airport"].tolist())
    _add_body(
        doc,
        f"""
        本项目把机场视为节点，把 Origin-Dest 航线视为有向边。边权不是单纯航班量，而是结合航班量、平均延误和预测风险。节点关键性则综合流量、介数中心性、平均预测风险和历史延误率。当前样本中综合关键性前五为 {top_airports}。

        该结果的管理含义是：某些机场的延误率未必最高，但它们连接的航线多、处于传播路径中心，受到同等冲击时更可能放大全网影响。因此策略评价必须把“节点位置”纳入，而不能只按历史延误率排序。
        """,
    )
    _add_formula(doc, "w_ij = alpha F_ij + beta R_ij + gamma D_ij", "w_ij 为航线边权；F_ij 表示航班量，R_ij 表示预测风险，D_ij 表示平均延误或延误强度。")
    _add_df_table(doc, nodes.sort_values("criticality", ascending=False)[["airport", "departures", "delay_rate", "avg_risk", "betweenness", "criticality"]], "表 2  关键机场指标节选", max_rows=8)
    _add_picture(doc, "fig_20_airport_network.png", "图 12  主要机场有向航线网络", "节点大小和颜色共同表达流量与关键性；DEN、DFW 等节点处于较重要位置，是默认冲击和优先恢复分析对象。")
    _add_picture(doc, "fig_21_airport_criticality.png", "图 13  关键机场综合指数排名", "该图直接服务于恢复资源排序：优先节点不只由出港量决定，也受网络中心性和风险叠加影响。")
    _add_picture(doc, "fig_22_network_risk_scatter.png", "图 14  中心性、风险与航班量关系", "散点图把三个维度合在一起，能直观看出“高风险但不一定高中心”和“高中心但风险中等”的节点差别。")

    doc.add_heading("7 ISM 结构分析：从根源因素到恢复时间", level=1)
    _add_body(
        doc,
        """
        ISM 用于解释延误因素之间的层级结构。本文不把 ISM 结果解释为严格因果，而把它作为结构化讨论框架：底层因素包括天气、容量、计划和缓冲；中间层包括前序晚到、机场拥堵、航线连接；上层表现为跨机场传播、恢复时间增加和旅客影响扩大。

        这种结构对报告和汇报很重要，因为它把技术模型与系统工程课程要求连接起来：预测模型解释“哪里风险高”，网络模型解释“为什么会传导”，传播仿真解释“冲击之后怎么演化”，评价模型解释“在目标冲突下如何选方案”。
        """,
    )
    _add_picture(doc, "fig_23_ism_adjacency.png", "图 15  ISM 邻接矩阵", "邻接矩阵来自因素关系判断，表达因素之间是否存在直接影响。")
    _add_picture(doc, "fig_24_ism_reachability.png", "图 16  ISM 可达矩阵", "可达矩阵体现直接与间接影响链，是层级划分的基础。")

    doc.add_heading("8 状态传播模型与扰动仿真", level=1)
    _add_body(
        doc,
        f"""
        传播模型以机场小时级平均正延误分钟数为状态变量。状态方程中，A 表示机场之间的延误传播，B u(t) 表示恢复控制，G w(t) 表示外部冲击。为了避免无意义的全连接回归，A 的估计带有网络掩码，只允许自身和有航线连接的上游机场进入。估计后谱半径为 {prop['spectral_radius_final']:.3f}，单步 MAE 为 {prop['mae']:.2f} 分钟。

        仿真情景包括常态、高峰、天气冲击和枢纽容量下降。策略包括基准、统一缓冲、关键枢纽优先和动态组合。所有策略在相同初始状态、相同冲击、相同预算口径下比较，避免把参数不一致误判为策略优劣。
        """,
    )
    _add_formula(doc, "x(t+1) = A x(t) + B u(t) + G w(t) + epsilon(t)", "x(t) 为机场状态向量，A 为传播矩阵，u(t) 为恢复资源控制，w(t) 为外部扰动，epsilon(t) 为误差项。")
    _add_code_block(
        doc,
        "代码片段 3  带网络掩码的传播矩阵估计",
        """
for airport in airports:
    allowed = [airport] + upstream_neighbors[airport]
    X_lag = state_df[[f"lag_{a}" for a in allowed]]
    y_next = state_df[f"next_{airport}"]
    ridge.fit(X_lag, y_next)
    A[airport, allowed] = ridge.coef_
A = stabilize_matrix(A, target_radius=0.82)
        """,
    )
    _add_picture(doc, "fig_27_propagation_matrix.png", "图 17  传播矩阵 A", "矩阵中非零项主要集中在有直接运行联系的机场对上；谱半径小于 1 保证情景内状态不会无限发散。")
    _add_picture(doc, "fig_29_recovery_curves.png", "图 18  天气冲击下性能恢复曲线", "动态组合策略恢复较快，但成本和复杂度也更高，因此不能只看曲线最低点。")
    _add_picture(doc, "fig_30_scenario_delay_compare.png", "图 19  不同情景累计延误对比", "多情景比较避免策略只在某一个案例上表现好；天气和枢纽容量下降是区分策略的重要压力测试。")
    _add_picture(doc, "fig_32_cost_delay_pareto.png", "图 20  成本-延误 Pareto 关系", "统一缓冲降低延误但成本较高，基准成本低但延误大；动态组合的优势来自恢复效果与成本之间的折中。")
    _add_code_block(
        doc,
        "代码片段 4  策略仿真循环",
        """
for hour in range(1, horizon + 1):
    control = policy.allocate(state=x, risk=risk_score, criticality=criticality)
    shock = scenario.external_shock(hour)
    x = A @ x + B @ control + G @ shock
    trajectory.append(score_state(x, control, hour))
        """,
    )

    doc.add_heading("9 多准则评价与策略选择", level=1)
    _add_body(
        doc,
        """
        策略评价不能只用单一指标。本文设置累计延误、平均延误、恢复时间、最低性能、恢复损失面积、传播范围、延误航班比例、取消代理、相对成本和复杂度等指标。成本型指标先正向化，再结合 AHP 主观权重和熵权客观离散度，最后计算 TOPSIS 接近度。

        主偏好下，恢复收益类指标权重较高，动态组合排名第一。风险型决策使用情景概率和损失矩阵计算期望损失，结果更偏向成本保守，因此基准策略排名第一。这个排名反转不是矛盾，而是说明不同管理偏好下“最优”的含义不同。
        """,
    )
    _add_formula(doc, "e_j = -k * sum_i p_ij ln(p_ij),    w_j = (1-e_j) / sum_j(1-e_j)", "熵权根据指标离散度分配客观权重，离散度越大，指标区分策略的能力越强。")
    _add_formula(doc, "C_i = D_i^- / (D_i^+ + D_i^-)", "TOPSIS 接近度 C_i 越大，表示方案越接近正理想解、远离负理想解。")
    _add_formula(doc, "E(L_s) = sum_c P(c) L_s,c", "风险型决策按情景概率 P(c) 与策略损失 L_s,c 计算期望损失。")
    _add_df_table(doc, rankings, "表 3  主偏好下策略综合排名", max_rows=4)
    _add_df_table(doc, weights[["indicator", "ahp_weight", "entropy_weight", "combined_weight"]], "表 4  指标权重节选", max_rows=10)
    _add_df_table(doc, risk, "表 5  风险型决策期望损失", max_rows=4)
    _add_picture(doc, "fig_34_topsis_score.png", "图 21  TOPSIS 综合得分", "dynamic_combo 得分最高但领先幅度有限，这提示结论需要配合敏感性和风险准则解释。")
    _add_picture(doc, "fig_35_fuzzy_grades.png", "图 22  模糊综合评价等级分布", "模糊评价帮助把连续指标转化为等级语言，便于答辩时解释策略质量差异。")
    _add_picture(doc, "fig_37_weight_sensitivity.png", "图 23  权重敏感性与排名反转", "当 AHP 恢复偏好占比低于较高阈值时，baseline 在成本保守设定下反超；这构成推荐边界。")
    _add_code_block(
        doc,
        "代码片段 5  AHP/熵权/TOPSIS 综合评价",
        """
combined_w = lambda_ahp * ahp_w + (1 - lambda_ahp) * entropy_w
norm = normalize_strategy_matrix(metrics, cost_columns=cost_like)
weighted = norm * combined_w
d_pos = distance(weighted, weighted.max(axis=0))
d_neg = distance(weighted, weighted.min(axis=0))
topsis_score = d_neg / (d_pos + d_neg)
        """,
    )

    doc.add_heading("10 静态 Web Demo 与部署", level=1)
    _add_body(
        doc,
        """
        为了让模型链路能够被现场检查，本项目将关键结果导出为静态 Web Demo。页面不依赖 Python 后端，所有数据来自 `web/assets/data/flightresilience-data.json`，可以通过 GitHub Pages 直接部署。静态站按“总览 - 数据 - 预测 - 网络 - 仿真 - 决策 - 方法”的路径组织，适合答辩时边讲边演示。

        相比原 Streamlit 原型，静态站的作用更偏“答辩证据面板”：它能让评委看到样本范围、风险预测输入、机场关键性、恢复曲线、TOPSIS 排名和权重敏感性是相互连接的，而不是分散的图表截图。
        """,
    )
    for file, caption, reading in [
        ("static_web_desktop.png", "图 24  静态 Web 首屏桌面视图", "首屏直接给出研究对象、样本范围、航班量、机场节点和主推荐策略，避免把 Demo 做成普通说明页。"),
        ("static_web_mobile.png", "图 25  静态 Web 移动端视图", "移动端采用紧凑 KPI 和单列布局，保证答辩现场换设备时不出现文字溢出。"),
        ("static_web_fullpage.png", "图 26  静态 Web 全页证据链", "全页截图展示从数据规律到综合决策的完整路径，说明 Demo 可以作为报告和 PPT 的交互补充。"),
    ]:
        _add_picture(doc, file, caption, reading, width=5.9, source_dir=SCREENSHOTS_DIR)
    _add_code_block(
        doc,
        "代码片段 6  静态 Web 数据导出",
        """
payload = {
    "kpi": read_json(DEMO_DIR / "kpi_summary.json"),
    "networkNodes": read_json(DEMO_DIR / "network_nodes.json"),
    "scenarioResults": records_from_parquet(DEMO_DIR / "scenario_results.parquet"),
    "strategyRankings": records_from_csv(DEMO_DIR / "strategy_rankings.csv"),
}
Path("web/assets/data/flightresilience-data.json").write_text(json.dumps(payload, ensure_ascii=False))
        """,
    )

    doc.add_heading("11 结论、局限与个人实践体会", level=1)
    _add_body(
        doc,
        """
        本项目的第一点结论是，航空延误应被理解为网络恢复问题，而不是孤立航班分类问题。预测模型提供风险输入，复杂网络提供传播位置，状态方程提供动态演化，综合评价提供有边界的策略建议。

        第二点结论是，高流量机场不必然是最高风险机场，最高延误率机场也不必然是最关键传播节点。真正有管理意义的指标是流量、中心性、预测风险和历史延误的组合。

        第三点结论是，dynamic_combo 在恢复/韧性优先的主偏好下表现最好，但不能写成绝对最优。风险型决策和权重敏感性已经显示 baseline 在成本保守条件下会反超。这个结论边界反而提高了研究可信度，因为现实调度本来就不是单目标优化。

        实践体会方面，我最明显的收获是：系统工程课程项目不能只堆算法，而要回答系统边界、目标冲突、证据链和可实施性。时间切分、防泄漏、统一预算口径和静态 Demo 部署看似是工程细节，但它们决定了结论是否能被复现和追问。后续若有真实机场容量、机组轮转和恢复成本数据，可以把相对情景参数进一步替换为运营约束优化模型。
        """,
    )
    _add_df_table(doc, metrics_by_scenario[["scenario", "strategy", "cumulative_delay", "recovery_time", "min_performance", "strategy_cost"]], "表 6  多情景策略表现节选", max_rows=12)

    doc.add_heading("附录：复现路径与文件索引", level=1)
    _add_body(
        doc,
        """
        A. 数据字典：`data/data_dictionary.csv`；数据审计：`data/data_audit.json`；数据清单：`data/data_manifest.csv`。
        B. 核心矩阵：`reports/tables/ism_adjacency_matrix.csv`、`ism_reachability_matrix.csv`、`propagation_matrix_A.csv`、`ahp_judgment_matrix.csv`。
        C. 核心脚本：`scripts/01_prepare_data.py` 至 `scripts/11_export_static_web_assets.py`。
        D. 静态站：`web/index.html`、`web/assets/js/app.js`、`web/assets/data/flightresilience-data.json`。
        E. 部署：推送到 `2711944586/xtgc` 后由 `.github/workflows/pages.yml` 发布 `web/`。
        """,
    )

    out = REPORTS_DIR / "report.docx"
    doc.save(out)
    return out


def _register_fonts() -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont("CNFont", str(path)))
                return "CNFont"
            except Exception:
                continue
    return "Helvetica"


def _pdf_table(df: pd.DataFrame, cols: list[str], max_rows: int = 6) -> Table:
    view = df[cols].head(max_rows).copy()
    data = [cols]
    for _, row in view.iterrows():
        data.append([f"{v:.3g}" if isinstance(v, float) else str(v) for v in row])
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#163A46")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CFDAD5")),
                ("FONTNAME", (0, 0), (-1, -1), "CNFont"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.2),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F6F4")]),
            ]
        )
    )
    return table


def _pdf_figure(path: Path, caption: str, max_width: float = 16.4 * cm) -> list:
    if not path.exists():
        return []
    img = Image(str(path), width=max_width, height=max_width * 0.58)
    return [Spacer(1, 0.15 * cm), img, Paragraph(caption, _pdf_styles()["Caption"])]


def _pdf_styles() -> dict:
    font = "CNFont"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CNTitle", fontName=font, fontSize=18, leading=24, alignment=1, textColor=colors.HexColor("#163A46"), spaceAfter=10))
    styles.add(ParagraphStyle(name="CNH1", fontName=font, fontSize=13, leading=18, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#163A46")))
    styles.add(ParagraphStyle(name="CNH2", fontName=font, fontSize=10.5, leading=15, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#147C7C")))
    styles.add(ParagraphStyle(name="CNBody", fontName=font, fontSize=8.8, leading=14, spaceAfter=5, textColor=colors.HexColor("#172326")))
    styles.add(ParagraphStyle(name="Formula", fontName=font, fontSize=8.2, leading=12, leftIndent=8, rightIndent=8, backColor=colors.HexColor("#F7E7D9"), textColor=colors.HexColor("#163A46"), spaceAfter=5))
    styles.add(ParagraphStyle(name="Caption", fontName=font, fontSize=7.4, leading=10, alignment=1, textColor=colors.HexColor("#6C7B80"), spaceAfter=5))
    return styles


def _pdf_section(story: list, title: str, paragraphs: Iterable[str], styles: dict) -> None:
    story.append(Paragraph(title, styles["CNH1"]))
    for para in paragraphs:
        story.append(Paragraph(para, styles["CNBody"]))


def generate_pdf_report() -> Path:
    ensure_dirs()
    _register_fonts()
    styles = _pdf_styles()
    audit = read_json(DATA_DIR / "data_audit.json")
    model = read_json(REPORTS_DIR.parent / "models" / "model_summary.json")
    prop = read_json(TABLES_DIR / "propagation_validation.json")
    decision = read_json(REPORTS_DIR.parent / "data" / "demo" / "decision_summary.json")
    nodes = _read_table("airport_nodes.csv")
    rankings = _read_table("strategy_rankings.csv")
    weights = _read_table("indicator_weights.csv")
    risk = _read_table("risk_decision.csv")

    out = REPORTS_DIR / "report.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=1.4 * cm, leftMargin=1.4 * cm, topMargin=1.4 * cm, bottomMargin=1.3 * cm)
    story: list = []
    story.append(Paragraph("面向突发扰动的航空网络延误传播识别与恢复策略决策", styles["CNTitle"]))
    story.append(Paragraph(f"样本：{audit['date_min']} 至 {audit['date_max']}，前 15 个机场，{audit['scoped_rows_top_airports']:,} 条记录；主推荐：{decision['recommended_strategy']}", styles["CNBody"]))

    _pdf_section(
        story,
        "摘要与问题背景",
        [
            "本文将航空延误从单航班分类问题扩展为网络恢复决策问题。研究链路覆盖真实数据、计划阶段预测、机场复杂网络、ISM 结构分析、状态空间仿真、多准则评价和静态 Web 演示。",
            f"最佳计划阶段模型测试 ROC-AUC 为 {model['best_test_metrics']['roc_auc']:.3f}，传播矩阵谱半径为 {prop['spectral_radius_final']:.3f}。主偏好下 dynamic_combo 排名第一，但风险/成本保守条件下 baseline 会反超。",
        ],
        styles,
    )
    story.append(Paragraph("p_i = P(y_i=1|x_i)=f(x_i)；x(t+1)=A x(t)+B u(t)+G w(t)+epsilon(t)；C_i=D_i^-/(D_i^+ + D_i^-)", styles["Formula"]))

    _pdf_section(
        story,
        "系统工程方法与数据防泄漏",
        [
            "霍尔三维结构组织问题界定、模型建立、方案优化和决策反馈；切克兰德软系统方法处理低成本、快恢复和公平可实施之间的价值冲突。",
            "训练、验证、测试采用时间顺序划分；历史统计特征只由训练窗口估计；实际起降、实际延误、延误原因和事后滑行信息不进入计划阶段模型。",
        ],
        styles,
    )
    story.append(Preformatted('route_rate = train.groupby("route")["ArrDel15"].mean()\nvalid["route_hist_delay_rate"] = valid["route"].map(route_rate).fillna(base_rate)', styles["CNBody"]))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_06_delay_distribution.png", "图 1  到达延误分钟数分布"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_07_daily_volume_delay.png", "图 2  每日航班量与延误率趋势"))
    story.append(PageBreak())
    story.extend(_pdf_figure(FIGURES_DIR / "fig_08_week_hour_heatmap.png", "图 3  星期-小时延误率热力图"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_09_volume_delay_scatter.png", "图 4  航班量与延误率关系"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_10_airport_delay_rank.png", "图 5  主要机场延误率排名"))

    story.append(PageBreak())
    _pdf_section(
        story,
        "预测、网络与结构解释",
        [
            "预测模型为后续传播仿真提供风险输入。SHAP 结果显示航线历史延误率、机场小时历史延误率和时间特征贡献较高。",
            "机场网络综合流量、介数中心性、平均预测风险和历史延误率识别关键节点。高流量机场不一定是最危险节点，关键在网络位置和风险叠加。",
        ],
        styles,
    )
    story.extend(_pdf_figure(FIGURES_DIR / "fig_14_model_metrics.png", "图 6  时间外测试集模型指标"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_15_confusion_matrix.png", "图 7  最佳模型混淆矩阵"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_17_calibration_curve.png", "图 8  概率校准曲线"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_18_shap_summary.png", "图 9  SHAP 全局重要性"))
    story.append(PageBreak())
    story.extend(_pdf_figure(FIGURES_DIR / "fig_20_airport_network.png", "图 10  主要机场网络"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_21_airport_criticality.png", "图 11  关键机场排名"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_22_network_risk_scatter.png", "图 12  中心性、风险与流量关系"))
    story.append(_pdf_table(nodes.sort_values("criticality", ascending=False), ["airport", "departures", "delay_rate", "avg_risk", "criticality"], 8))
    story.append(PageBreak())
    story.extend(_pdf_figure(FIGURES_DIR / "fig_25_ism_hierarchy.png", "图 13  ISM 递阶结构"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_23_ism_adjacency.png", "图 14  ISM 邻接矩阵"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_24_ism_reachability.png", "图 15  ISM 可达矩阵"))

    story.append(PageBreak())
    _pdf_section(
        story,
        "传播仿真与策略评价",
        [
            "状态传播模型采用带网络掩码的 Ridge 滞后回归估计传播矩阵 A，并通过谱半径约束保持仿真稳定。四种策略在相同初始状态、相同冲击和相同预算口径下比较。",
            "综合评价阶段使用 AHP/熵权组合权重、TOPSIS 接近度、模糊综合评价和风险型决策。dynamic_combo 是恢复优先偏好下的推荐，不是无条件最优。",
        ],
        styles,
    )
    story.append(Paragraph("w_ij=alpha F_ij+beta R_ij+gamma D_ij；E(L_s)=sum_c P(c)L_s,c", styles["Formula"]))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_27_propagation_matrix.png", "图 16  传播矩阵 A"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_29_recovery_curves.png", "图 17  天气冲击下恢复曲线"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_30_scenario_delay_compare.png", "图 18  多情景累计延误对比"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_32_cost_delay_pareto.png", "图 19  成本-延误 Pareto 关系"))
    story.append(_pdf_table(rankings, ["strategy", "topsis_score", "cumulative_delay", "recovery_time", "min_performance", "strategy_cost"], 4))
    story.append(_pdf_table(weights, ["indicator", "ahp_weight", "entropy_weight", "combined_weight"], 8))
    story.append(_pdf_table(risk, ["strategy", "expected_loss", "risk_rank"], 4))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_34_topsis_score.png", "图 20  TOPSIS 综合得分"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_35_fuzzy_grades.png", "图 21  模糊综合评价等级分布"))
    story.extend(_pdf_figure(FIGURES_DIR / "fig_37_weight_sensitivity.png", "图 22  权重敏感性"))

    story.append(PageBreak())
    _pdf_section(
        story,
        "静态 Web Demo、结论与体会",
        [
            "静态 Web Demo 由 `web/index.html`、`assets/js/app.js` 和预计算 JSON 驱动，可通过 GitHub Pages 部署到 2711944586/xtgc。页面把数据规律、风险预测、机场网络、扰动仿真、综合评价串成可交互演示线。",
            "个人体会是：系统工程项目不能只堆模型，而要把系统边界、目标冲突、技术步骤、结果边界和可复现交付统一起来。防泄漏、统一预算口径和静态部署这些工程细节，直接决定研究结论能否经得住追问。",
        ],
        styles,
    )
    story.extend(_pdf_figure(SCREENSHOTS_DIR / "static_web_desktop.png", "图 23  静态 Web 首屏"))
    story.extend(_pdf_figure(SCREENSHOTS_DIR / "static_web_mobile.png", "图 24  静态 Web 移动端视图", max_width=9.0 * cm))
    story.extend(_pdf_figure(SCREENSHOTS_DIR / "static_web_fullpage.png", "图 25  静态 Web 全页证据链", max_width=15.4 * cm))
    story.append(Paragraph("附录：完整字段字典、清洗记录、ISM 矩阵、传播矩阵、AHP 判断矩阵、策略长表和静态站数据均保存在仓库 data/、reports/tables/、web/ 目录。", styles["CNBody"]))
    doc.build(story)
    return out


def generate_script() -> Path:
    metrics = _read_table("strategy_rankings.csv")
    top = metrics.iloc[0]
    script = f"""# FlightResilience 7分30秒讲稿

> 使用方式：每页先讲“核心句”，时间紧时只讲“可删减句”之前的内容。Demo 如果现场加载异常，就用 PPT 中的静态截图继续讲，不影响主线。

## 第 1 页，25 秒
核心句：我们研究的不是某个航班会不会延误，而是一个局部扰动如何演化成航空网络恢复问题。

详细稿：大家好，我们的题目是“面向突发扰动的航空网络延误传播识别与恢复策略决策”。一个航班晚点看起来是局部事件，但在枢纽化运行中，飞机、机组和旅客连接会把局部延误传到其他机场。我们希望回答三个问题：延误会传到哪里，哪些机场应优先恢复，有限资源下哪种策略更合适。

过渡句：所以第一页先把问题从“预测”推到“网络恢复”。

可删减句：如果只看单航班准确率，就会漏掉传播路径和恢复资源配置。

Demo 兜底：如果网页打不开，直接说明 PPT 的网络图和恢复曲线就是 Web Demo 的核心输出。

## 第 2 页，40 秒
核心句：单航班预测不足以支撑恢复决策，因为它不回答传播和资源分配。

详细稿：传统延误预测输出概率或标签，比如某一航班是否会延误 15 分钟。但管理者真正关心的是冲击后的系统表现：延误是否会扩散，哪个节点最值得优先处理，缓冲资源应该均匀投放还是投向关键枢纽。因此我们把研究对象定义为包含机场、航线、航班计划、航空公司和恢复资源的航空运行系统。

过渡句：有了这个定义，下一步就是明确系统边界和目标冲突。

可删减句：预测是输入，不是最终答案。

## 第 3 页，40 秒
核心句：系统边界让多主体目标冲突变得可讨论。

详细稿：系统内部包括机场容量、航线连接、航班计划、运行状态和恢复资源；外部环境包括天气、高峰需求、空域限制和枢纽容量下降。乘客希望少延误，航空公司关注成本，机场关注容量恢复，监管者关注安全与公平，所以评价指标不能只有一个。我们最终设置了效率、韧性、航班影响、成本和复杂度等指标。

过渡句：这也是为什么我们采用系统工程方法链，而不是单模型路线。

可删减句：多目标冲突是本项目区别于普通预测题的关键。

## 第 4 页，45 秒
核心句：方法链把结构解释、预测模型、传播仿真和综合评价连成闭环。

详细稿：我们用霍尔三维结构组织流程：先明确问题，再做系统设计、方案综合、建模、优化和评价。切克兰德软系统方法用于处理不同主体的价值冲突。结构层用鱼骨图和 ISM，数据层用机器学习和 SHAP，网络层用中心性和关键性指标，动态层用状态空间传播模型，决策层用 AHP、熵权、TOPSIS、模糊评价和风险决策。

过渡句：下面先看数据本身揭示了什么。

可删减句：这一页要强调“方法之间有接口”，不是把算法平铺。

## 第 5 页，45 秒
核心句：BTS 数据显示，延误在时间和机场之间高度不均匀。

详细稿：数据来自美国交通部 BTS 航班准点数据，样本是 2024 年 1 到 3 月、前 15 个主要机场，共 224755 条记录。热力图显示不同时段延误率差异明显，机场排名也说明延误不是平均分布。散点图进一步说明航班量大不等于延误率最高，所以恢复策略不能只按流量排队。

过渡句：既然延误有规律，我们先建立计划阶段风险输入。

可删减句：这里的重点是“异质性”，不是只报样本量。

## 第 6 页，45 秒
核心句：ISM 解释了延误因素如何从根源条件传导到恢复时间。

详细稿：我们先通过鱼骨图整理天气、容量、计划、缓冲、前序延误和信息协同等因素，再构造 ISM 递阶结构。底层因素通过机场拥堵和上游晚到向上传导，最后表现为跨机场传播和恢复时间增加。这里我们不把 ISM 说成严格因果，而把它作为系统结构解释框架。

过渡句：结构解释之后，再看计划阶段能否识别风险。

可删减句：主动说明“不夸大因果”会让结论更可信。

## 第 7 页，50 秒
核心句：预测模型的价值是提供风险输入，而不是单独替代调度。

详细稿：预测目标是 ArrDel15，也就是到达延误是否超过 15 分钟。我们采用时间顺序划分，避免随机切分造成信息泄漏；实际起降时间、实际延误分钟数和延误原因都不进入模型。最终随机森林在时间外测试集上的 ROC-AUC 约为 0.675，PR-AUC 约为 0.391。SHAP 显示航线历史延误率、机场小时历史延误率和时间特征贡献较高。

过渡句：有了风险输入，下一步就要看风险在网络里的位置。

可删减句：这里不要把模型指标讲成特别夸张，它是计划阶段的可用输入。

## 第 8 页，50 秒
核心句：关键机场不是按航班量简单排序，而是由流量、风险和网络位置共同决定。

详细稿：我们把机场作为节点、航线作为有向边，计算介数中心性、PageRank、预测风险和历史延误率，再合成综合关键性。结果中 DEN、DFW、LAX、MCO 和 LAS 排名靠前。这个结果说明，同样的延误发生在不同机场，传播后果不一样；高中心性节点更可能把局部冲击放大为全网恢复问题。

过渡句：识别关键节点之后，就可以做扰动传播仿真。

可删减句：如果时间紧，只强调“高流量不等于最高关键性”。

## 第 9 页，60 秒
核心句：状态空间模型让四种恢复策略可以在同一冲击下公平比较。

详细稿：我们把机场小时级平均正延误作为状态，用方程 x(t+1)=A x(t)+B u(t)+G w(t)+epsilon(t) 表示传播。A 是带网络掩码估计的传播矩阵，只允许自身和有航线连接的上游机场影响当前机场。估计后谱半径约 0.820，单步 MAE 约 11.63 分钟。策略包括基准、统一缓冲、关键枢纽优先和动态组合，全部在相同初始状态、相同冲击和统一预算口径下比较。

过渡句：接下来进入 Demo，直接看策略曲线怎么变。

可删减句：谱半径小于 1 的意义是仿真不会无界发散。

## 第 10 页，90 秒
核心句：Demo 的演示路径是从数据规律走到策略排名，而不是单独展示一个页面。

详细稿：打开静态 Web 后，先看总览的样本范围、航班量和主推荐。然后在数据页选择 DEN，观察热力图和机场风险排名；在预测页调整小时、距离和拥堵水平，看到风险概率变化；在网络页点击机场节点，解释关键性；在仿真页选择天气或枢纽容量下降情景，切换四种策略。可以看到动态组合通常恢复更快，当前主结果中排名第一策略的平均累计延误约为 {top['cumulative_delay']:.0f}，平均恢复时间约 {top['recovery_time']:.1f} 小时。

过渡句：但策略恢复快，不代表无条件最优，还要看偏好和风险。

可删减句：如果现场只演示一个页面，就演示仿真页和决策页。

Demo 兜底：网页若受网络或浏览器限制，展示 PPT 中的静态 Web 截图，并按同一路径讲。

## 第 11 页，55 秒
核心句：推荐是有条件的，取决于管理偏好、预算和风险态度。

详细稿：综合 AHP 主观偏好和熵权客观信息后，TOPSIS 主排序推荐 dynamic_combo，因为它在恢复速度和系统性能上表现好。但敏感性分析显示，当 λ 较低、也就是更强调客观成本和保守准则时，baseline 可能反超。风险型决策按情景概率和损失矩阵计算，也会偏向低成本方案。因此我们的表述是“恢复/韧性优先下推荐动态组合”，不是绝对最优。

过渡句：最后总结项目贡献和个人体会。

可删减句：排名反转不是失败，而是说明模型能暴露管理取舍。

## 第 12 页，45 秒
核心句：本项目的贡献是把系统工程闭环落到可运行、可复现、可演示的成果。

详细稿：第一，问题上把延误从单航班预测扩展到网络恢复；第二，技术上把真实数据、预测、网络、ISM、传播仿真和多准则评价连接起来；第三，交付上提供报告、PPT、讲稿和可静态部署 Web Demo。局限也很明确：公开数据缺少真实企业恢复成本、机组轮转和容量约束，因此成本与资源参数是透明的相对情景设定。我的实践体会是，系统工程项目最重要的不是堆模型，而是把系统边界、目标冲突、证据链和结论边界讲清楚。

结束句：以上就是我们的汇报，谢谢大家。

可删减句：如果超时，最后只保留三点贡献和一个局限。
"""
    out = REPORTS_DIR.parent / "slides" / "script.md"
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(script, encoding="utf-8")
    return out


def generate_deliverables() -> dict[str, str]:
    docx = generate_docx_report()
    pdf = generate_pdf_report()
    script = generate_script()
    return {"docx": str(docx), "pdf": str(pdf), "script": str(script)}
