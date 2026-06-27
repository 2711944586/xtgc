from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "提交包" / "航空网络延误传播与恢复策略报告.docx"
CANONICAL_OUT = ROOT / "提交包" / "01_航空网络延误传播与恢复策略_报告.docx"
FIG = ROOT / "提交包" / "06_补充图表与矩阵" / "figures"
SCREEN = ROOT / "提交包" / "07_Web截图"
TABLES = ROOT / "reports" / "tables"
DATA = ROOT / "data"
DEMO = ROOT / "data" / "demo"
MODELS = ROOT / "models"
ASSETS = ROOT / "tmp" / "report_assets"

FONT_CN = "Microsoft YaHei"
FONT_WEST = "Calibri"
FONT_MONO = "Consolas"

NAVY = RGBColor(20, 48, 68)
TEAL = RGBColor(20, 113, 112)
INK = RGBColor(31, 43, 46)
MUTED = RGBColor(91, 105, 112)
GOLD = RGBColor(176, 111, 35)
WHITE = RGBColor(255, 255, 255)

FILL_SOFT = "EEF5F2"
FILL_BLUE = "143044"
FILL_TEAL = "147170"
FILL_GOLD = "F7E8D6"
FILL_GRAY = "F3F6F8"
LINE = "C9D8D4"


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def hex_rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(
    run,
    *,
    name: str = FONT_CN,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts()
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_WEST if name == FONT_CN else name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_WEST if name == FONT_CN else name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=4, line=1.12) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_rule(paragraph, color: str = LINE, size: str = "8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run("1")
    set_run_font(run, size=8.5, color=MUTED)

    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: Sequence[int], *, indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths_dxa)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color: str = LINE, size: str = "4") -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def clear_cell(cell) -> None:
    cell.text = ""
    for p in cell.paragraphs:
        p.text = ""


def add_cell_text(
    cell,
    text: str,
    *,
    size: float = 8.6,
    color: RGBColor = INK,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, after=0, line=1.08)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    size: float = 9.4,
    color: RGBColor = INK,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 4,
    line: float = 1.14,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> object:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_body(doc: Document, text: str, *, size: float = 9.35) -> None:
    for raw in text.strip().split("\n\n"):
        para = " ".join(line.strip() for line in raw.splitlines() if line.strip())
        if para:
            add_para(doc, para, size=size, after=4, line=1.15)


def add_labeled_points(
    doc: Document,
    title: str,
    items: Sequence[tuple[str, str]],
    *,
    size: float = 8.9,
) -> None:
    add_para(doc, title, size=9.0, color=TEAL, bold=True, before=2, after=2)
    for label, text in items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.12)
        r = p.add_run(f"{label}：")
        set_run_font(r, size=size, color=NAVY, bold=True)
        r = p.add_run(text)
        set_run_font(r, size=size, color=INK)


def add_bullet(doc: Document, items: Iterable[str], *, size: float = 9.1) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=2, line=1.12)
        run = p.add_run(item)
        set_run_font(run, size=size, color=INK)


def add_numbered(doc: Document, items: Iterable[str], *, size: float = 9.1) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=2, line=1.12)
        run = p.add_run(item)
        set_run_font(run, size=size, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        paragraph_rule(p, color="D3E0DD", size="6")


def add_callout(doc: Document, title: str, text: str, *, fill: str = FILL_SOFT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9000], indent_dxa=120)
    set_table_borders(table, color="D5E1DE", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.12)
    r = p.add_run(title)
    set_run_font(r, size=9.2, color=TEAL, bold=True)
    p.add_run("\n")
    r = p.add_run(text)
    set_run_font(r, size=8.9, color=INK)
    doc.add_paragraph()


def add_formula(doc: Document, formula: str, note: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9000], indent_dxa=120)
    set_table_borders(table, color="E6D0B3", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL_GOLD)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.08)
    r = p.add_run(formula)
    set_run_font(r, name=FONT_MONO, size=8.6, color=NAVY, bold=True)
    p.add_run("\n")
    r = p.add_run(note)
    set_run_font(r, size=8.6, color=INK)
    doc.add_paragraph()


def add_table_from_rows(
    doc: Document,
    caption: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths_dxa: Sequence[int],
    *,
    max_rows: int | None = None,
) -> None:
    add_para(doc, caption, size=8.8, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, FILL_BLUE)
        add_cell_text(cell, header, size=8.0, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    limited = rows if max_rows is None else rows[:max_rows]
    for ridx, row_values in enumerate(limited):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if isinstance(value, float):
                if abs(value) < 1:
                    text = f"{value:.3f}"
                elif abs(value) < 100:
                    text = f"{value:.2f}"
                else:
                    text = f"{value:,.0f}"
            else:
                text = str(value)
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or len(str(text)) < 12 else WD_ALIGN_PARAGRAPH.LEFT
            add_cell_text(cells[idx], text, size=7.7, color=INK, align=align)
    doc.add_paragraph()


def add_compact_matrix(
    doc: Document,
    caption: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths_dxa: Sequence[int],
    *,
    header_fill: str = FILL_TEAL,
    body_size: float = 7.3,
    header_size: float = 7.5,
) -> None:
    add_para(doc, caption, size=8.5, color=MUTED, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths_dxa, indent_dxa=120)
    set_table_borders(table, color="D7E2DF", size="4")
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        add_cell_text(cell, str(header), size=header_size, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ridx, row_values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            text = str(value)
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 or len(text) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            add_cell_text(cells[idx], text, size=body_size, color=INK, align=align)
    doc.add_paragraph()


def add_figure_cell(cell, image_path: Path, caption: str, reading: str, *, width_in: float = 2.95) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=1, line=1.0)
    if image_path.exists():
        p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, after=1, line=1.05)
    r = cap.add_run(caption)
    set_run_font(r, size=7.8, color=MUTED, bold=True)
    exp = cell.add_paragraph()
    set_paragraph_spacing(exp, after=0, line=1.08)
    r = exp.add_run("图意：")
    set_run_font(r, size=7.7, color=TEAL, bold=True)
    r = exp.add_run(reading)
    set_run_font(r, size=7.7, color=INK)


def add_picture_grid(
    doc: Document,
    figures: Sequence[tuple[str, str, str]],
    *,
    source: Path = FIG,
    width_in: float = 2.95,
) -> None:
    for i in range(0, len(figures), 2):
        pair = figures[i : i + 2]
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, [4500, 4500], indent_dxa=120)
        set_table_borders(table, color="D7E2DF", size="4")
        for idx, (file, caption, reading) in enumerate(pair):
            cell = table.cell(0, idx)
            set_cell_shading(cell, "FFFFFF")
            add_figure_cell(cell, source / file, caption, reading, width_in=width_in)
        if len(pair) == 1:
            set_cell_shading(table.cell(0, 1), FILL_GRAY)
            add_cell_text(table.cell(0, 1), "本栏留白用于保持图文组平衡。", size=8, color=MUTED)
        doc.add_paragraph()


def add_full_picture(doc: Document, file: str, caption: str, reading: str, *, source: Path = FIG, width=6.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2, line=1.0)
    path = source / file
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(width))
    add_para(doc, caption, size=8.5, color=MUTED, bold=True, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, "读图解释", reading, fill="F8FBFA")


def fit_image_to_canvas(src: Path, dst: Path, *, canvas=(1800, 1120), padding=70) -> Path:
    from PIL import Image

    dst.parent.mkdir(parents=True, exist_ok=True)
    image = Image.open(src).convert("RGB")
    max_w = canvas[0] - padding * 2
    max_h = canvas[1] - padding * 2
    image.thumbnail((max_w, max_h), Image.LANCZOS)
    out = Image.new("RGB", canvas, "white")
    x = (canvas[0] - image.width) // 2
    y = (canvas[1] - image.height) // 2
    out.paste(image, (x, y))
    fmt = "JPEG" if dst.suffix.lower() in {".jpg", ".jpeg"} else "PNG"
    if fmt == "JPEG":
        out.save(dst, format=fmt, quality=92, optimize=True)
    else:
        out.save(dst, format=fmt)
    return dst


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_WEST)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_WEST)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.14

    for name, size, color, before, after in [
        ("Heading 1", 14.2, NAVY, 12, 6),
        ("Heading 2", 11.4, TEAL, 8, 4),
        ("Heading 3", 10.2, GOLD, 6, 3),
    ]:
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_WEST)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_WEST)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(9.0)
        style.paragraph_format.left_indent = Inches(0.22)
        style.paragraph_format.first_line_indent = Inches(-0.12)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("FlightResilience | 航空网络延误传播与恢复策略系统工程报告")
    set_run_font(r, size=8.1, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("数据：U.S. DOT BTS TranStats 2024Q1 | Page ")
    set_run_font(r, size=8.0, color=MUTED)
    add_field(footer, "PAGE")
    r = footer.add_run(" / ")
    set_run_font(r, size=8.0, color=MUTED)
    add_field(footer, "NUMPAGES")


def add_cover(doc: Document, audit: dict, kpi: dict, decision: dict, prop: dict, model: dict) -> None:
    add_para(doc, "SYSTEM ENGINEERING COURSE PROJECT", size=8.8, color=TEAL, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "面向突发扰动的航空网络延误传播识别", size=22.0, color=NAVY, bold=True, after=0, line=1.02, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "与恢复策略决策", size=21.0, color=NAVY, bold=True, after=6, line=1.02, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "基于 BTS 公开航班数据的系统工程方法链、传播仿真与多准则决策报告", size=11.5, color=MUTED, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, [1800, 7200], indent_dxa=120)
    set_table_borders(meta, color="D6E2DF", size="4")
    rows = [
        ("研究对象", "美国主要机场网络在天气、高峰需求、枢纽容量下降等扰动下的延误传播与恢复策略选择"),
        ("系统边界", "机场节点、航线边、航班计划、历史延误状态、预测风险、恢复资源与策略成本"),
        ("课程方法", "霍尔三维结构、切克兰德软系统方法、系统分析 5W1H、解释结构模型化技术(ISM)、状态空间模型、AHP（层次分析法）/熵权/TOPSIS、模糊综合评价、风险型与不确定型决策"),
        ("交付内容", "Word 报告、PDF 报告、PPT、讲稿、静态 Web Demo、补充图表与矩阵、核心代码"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = meta.rows[idx].cells
        set_cell_shading(cells[0], FILL_BLUE)
        add_cell_text(cells[0], label, size=8.6, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(cells[1], value, size=8.5, color=INK)

    doc.add_paragraph()
    strip = doc.add_table(rows=2, cols=5)
    strip.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(strip, [1800, 1800, 1800, 1800, 1800], indent_dxa=120)
    set_table_borders(strip, color="D6E2DF", size="4")
    metrics = [
        ("样本航班", f"{kpi['flights']:,}"),
        ("机场节点", f"{kpi['airport_count']}"),
        ("延误率", f"{kpi['delay_rate']:.1%}"),
        ("模型 AUC", f"{model['best_test_metrics']['roc_auc']:.3f}"),
        ("谱半径", f"{prop['spectral_radius_final']:.3f}"),
    ]
    for idx, (label, value) in enumerate(metrics):
        set_cell_shading(strip.rows[0].cells[idx], FILL_TEAL)
        add_cell_text(strip.rows[0].cells[idx], label, size=7.8, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(strip.rows[1].cells[idx], FILL_SOFT)
        add_cell_text(strip.rows[1].cells[idx], value, size=12, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(
        doc,
        f"样本范围：{audit['date_min']} 至 {audit['date_max']}；主偏好推荐：{format_strategy(decision['recommended_strategy'])}；风险/不确定决策下推荐：{format_strategy(decision['risk_recommended_strategy'])}。",
        size=8.8,
        color=MUTED,
        after=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    p = add_para(doc, "课程：系统工程    项目：FlightResilience    文档版本：升级版 Word 报告", size=8.7, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph_rule(p, color="D4E0DD", size="8")
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_para(doc, "目录", size=18, color=NAVY, bold=True, after=6)
    add_para(doc, "目录按三级标题自动生成，用于快速定位方法、图表、仿真、评价与结论。", size=8.7, color=MUTED, after=6)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def format_strategy(value: str) -> str:
    return {
        "baseline": "基准策略",
        "uniform_buffer": "统一缓冲",
        "hub_priority": "关键枢纽优先",
        "dynamic_combo": "动态组合",
    }.get(value, value)


def strategy_rows(df: pd.DataFrame) -> list[list[object]]:
    rows: list[list[object]] = []
    for _, row in df.iterrows():
        rows.append(
            [
                format_strategy(row["strategy"]),
                row["topsis_score"],
                row["cumulative_delay"],
                row["recovery_time"],
                row["min_performance"],
                row["strategy_cost"],
                int(row["overall_rank"]),
            ]
        )
    return rows


def refresh_with_word(path: Path) -> None:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        print(f"Word COM unavailable: {exc!r}")
        return
    work = Path(tempfile.mkdtemp(prefix="fr_word_refresh_"))
    input_docx = work / "input.docx"
    refreshed_docx = work / "refreshed.docx"
    shutil.copy2(path, input_docx)
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        try:
            word.AutomationSecurity = 3
        except Exception:
            pass
        doc = word.Documents.Open(str(input_docx.resolve()), False, False, False)
        try:
            doc.Fields.Update()
            for i in range(1, doc.TablesOfContents.Count + 1):
                doc.TablesOfContents(i).Update()
            doc.SaveAs2(str(refreshed_docx.resolve()), FileFormat=16)
        finally:
            doc.Close(False)
    finally:
        word.Quit()
    shutil.copy2(refreshed_docx, path)


def export_pdf_with_word(docx: Path, pdf: Path) -> None:
    import win32com.client  # type: ignore

    pdf.parent.mkdir(parents=True, exist_ok=True)
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(docx.resolve()), False, True, False)
        try:
            doc.ExportAsFixedFormat(str(pdf.resolve()), 17)
        finally:
            doc.Close(False)
    finally:
        word.Quit()


def build_report() -> Path:
    audit = read_json(DATA / "data_audit.json")
    kpi = read_json(DEMO / "kpi_summary.json")
    decision = read_json(DEMO / "decision_summary.json")
    model = read_json(MODELS / "model_summary.json")
    prop = read_json(TABLES / "propagation_validation.json")
    network_summary = read_json(TABLES / "network_summary.json")

    model_metrics = pd.read_csv(TABLES / "model_metrics.csv")
    nodes = pd.read_csv(TABLES / "airport_nodes.csv")
    rankings = pd.read_csv(TABLES / "strategy_rankings.csv")
    weights = pd.read_csv(TABLES / "indicator_weights.csv")
    risk = pd.read_csv(TABLES / "risk_decision.csv")
    uncertainty = pd.read_csv(TABLES / "uncertainty_decision.csv")
    fuzzy = pd.read_csv(TABLES / "fuzzy_evaluation.csv")
    scenarios = pd.read_csv(TABLES / "strategy_metrics_by_scenario.csv")
    manifest = pd.read_csv(DATA / "data_manifest.csv")

    doc = Document()
    style_document(doc)
    add_cover(doc, audit, kpi, decision, prop, model)
    add_toc(doc)

    add_heading(doc, "摘要与核心结论", 1)
    add_body(
        doc,
        f"""
        本报告围绕“突发扰动下航空网络延误如何传播、关键机场如何识别、恢复资源如何配置”展开。与只做单点航班延误预测不同，本文把机场、航线、航班计划、历史运行状态、外部扰动和恢复策略放入同一个系统边界中，形成“数据审计 - 风险预测 - 网络建模 - 结构解释 - 状态传播 - 多准则评价 - Web 交互展示”的闭环。

        数据层面，项目使用 U.S. DOT BTS TranStats 2024 年 1 月至 3 月公开航班准点数据，清洗后聚焦样本内总流量前 {kpi['airport_count']} 个机场，共 {kpi['flights']:,} 条航班记录，整体到达延误率为 {kpi['delay_rate']:.1%}。模型层面，随机森林在时间外测试集上取得 ROC-AUC {model['best_test_metrics']['roc_auc']:.3f}、PR-AUC {model['best_test_metrics']['pr_auc']:.3f}，其意义不是给出绝对准确的“是否延误”标签，而是为传播矩阵和恢复优先级提供可比较的计划阶段风险输入。

        网络层面，前五个综合关键机场为 {'、'.join(network_summary['top_critical_airports'])}。它们并非都只是流量最大机场，而是在流量、风险、中心性和历史延误叠加后更容易放大全网影响的节点。传播层面，带网络掩码的状态空间模型估计后谱半径为 {prop['spectral_radius_final']:.3f}，单步 MAE 为 {prop['mae']:.2f} 分钟，说明在设定情景中状态递推稳定且优于简单保持基线。

        决策层面，主偏好权重下 {format_strategy(decision['recommended_strategy'])} 的 TOPSIS 接近度最高；但风险型决策、Hurwicz 折中准则和最小最大后悔值等不确定型准则均提示，在成本极端保守或风险规避偏好下，{format_strategy(decision['risk_recommended_strategy'])} 会反超。因此本文给出的是“有条件推荐”：当目标偏向恢复效率与网络韧性时采用动态组合；当资源极紧、成本压倒性重要时保留基准或低强度策略。

        从写作结构看，报告采用“方法先行、证据跟随、结论回扣”的方式组织。每一组图片都对应一个系统工程问题：数据图回答系统状态是否存在异质性，模型图回答风险输入是否可信，网络图回答延误传播依赖哪些节点和边，仿真图回答扰动后系统如何演化，评价图回答多目标冲突下怎样选择策略。这样能让文字与图片相互解释，而不是图片堆在正文后面。
        """,
    )
    add_callout(
        doc,
        "报告升级重点",
        "新版报告把课程 PDF/参考材料中的霍尔三维结构、切克兰德软系统方法、解释结构模型化技术(ISM)、状态空间建模、系统评价、AHP（层次分析法）、模糊综合评价、风险型与不确定型决策逐一映射到项目环节，并让每组图表都配有“为什么做、怎么看、对决策意味着什么”的解释。",
    )
    add_body(
        doc,
        """
        为了让文字内容与图片真正对应，本报告采用三层读图口径：第一层说明图中直接看到的现象，例如延误是否长尾、机场是否集中、策略曲线是否恢复更快；第二层说明该现象在系统工程中的含义，例如整体性、关联性、动态性或可行性如何体现；第三层再回到决策，解释它怎样影响关键机场识别、传播仿真或策略排序。这样处理后，图片不只是视觉材料，而是系统分析、模型化和决策评价的证据节点。

        课程 PDF 中反复强调系统工程不是单纯求一个局部最优，而是在复杂系统中协调多目标、多主体和多阶段过程。本文据此把“预测准确率”“网络中心性”“恢复曲线”“TOPSIS 得分”“风险型/不确定型决策结果”放进同一条证据链：前面的数据图定义问题，网络和 ISM 图解释结构，状态空间图解释动态，评价和决策图给出可执行但有边界的建议。
        """,
    )
    add_compact_matrix(
        doc,
        "图表阅读总口径：从现象到方法再到结论",
        ["证据类型", "直接阅读", "系统工程含义", "决策用途"],
        [
            ["数据图", "延误率、航班量、时段热力和长尾分布。", "识别系统状态的时间异质性和空间异质性。", "决定预测模型和情景仿真的输入边界。"],
            ["网络图", "机场节点、航线边、中心性和风险分布。", "体现系统的关联性、层次性和有序结构。", "确定关键机场和枢纽优先策略的依据。"],
            ["ISM/矩阵图", "邻接关系、可达关系和递阶结构。", "把复杂因素转化为结构模型，连接课程模型化方法。", "说明为什么后续要建状态空间和多指标评价。"],
            ["仿真图", "恢复曲线、热力图、传播范围、Pareto 关系。", "刻画动态性、控制变量和外部扰动共同作用。", "生成策略评价矩阵和恢复优先级。"],
            ["评价图", "权重、TOPSIS、模糊等级、敏感性和风险对照。", "体现目标优化、可行性和不同决策态度。", "形成动态组合与基准策略的条件式推荐。"],
        ],
        [1500, 2500, 2700, 2600],
    )
    add_table_from_rows(
        doc,
        "课程 PDF 使用清单与本文对应位置",
        ["课程 PDF", "核心方法点", "本文使用方式"],
        [
            ["第一章 系统工程概述", "系统工程服务于复杂系统的预测、规划、评价、决策与管理，强调反馈控制和整体最优。", "将航空延误定义为机场网络运行系统问题，而不是孤立的航班分类问题；用反馈闭环解释模型、仿真、评价与 Web 交付。"],
            ["第二章 系统工程方法论", "整体性、有序相关性、动态性、目标优化、可行性原则；霍尔三维结构；切克兰德软系统方法；系统分析 5W1H。", "构建时间维、逻辑维、知识维三维方法链，用 5W1H 划定评价问题，并把多主体目标冲突写入指标体系。"],
            ["第三章 系统模型与模型化", "模型是现实系统的抽象；结构模型可用集合、有向图、邻接矩阵、可达矩阵表达；状态空间描述动态系统。", "建立风险预测模型、机场有向网络、ISM 矩阵与状态空间传播方程，形成可演示的思想试验。"],
            ["第四章 系统评价方法", "系统评价由 5W1H 构成；AHP（层次分析法）将复杂评价分层并两两比较；模糊综合评价处理等级性判断。", "设定评价对象、主体、目的、时期、地点与方法，融合 AHP（层次分析法）、熵权、TOPSIS 和模糊等级解释。"],
            ["第五章 系统决策分析", "决策包括准备、分析、选择、实施反馈；风险型决策用期望值；不确定型决策用悲观准则、乐观准则、后悔值法(Savage准则/最小最大后悔值准则)、折中准则、等可能性(Laplace准则)。", "在 TOPSIS 主推荐之外，补充期望损失、Laplace、Hurwicz、最小最大后悔值准则等结果，给出有条件推荐。"],
        ],
        [2100, 3500, 4300],
    )
    add_labeled_points(
        doc,
        "全文阅读路径与证据链",
        [
            ("先界定系统", "第 1 章说明研究对象、系统边界、输入输出和多主体目标冲突，对应系统工程概述中“整体性、关联性、环境适应性”的思想。"),
            ("再建立模型", "第 3-8 章依次完成数据审计、预测模型、机场网络、ISM 层级和状态空间传播模型，对应系统模型化中“用模型替代现实系统进行分析”的要求。"),
            ("最后评价决策", "第 9 章把仿真结果转化为多指标评价矩阵，并同时使用 TOPSIS、模糊综合评价、风险型决策和不确定型决策，避免只凭单张图或单个指标下结论。"),
            ("贯穿反馈", "第 10-11 章用 Web Demo、复现脚本、局限分析和后续改进形成反馈闭环，对应系统工程工作过程中的实施、检验和修订。"),
        ],
    )

    add_heading(doc, "1 问题界定、系统边界与目标冲突", 1)
    add_heading(doc, "1.1 从单点延误预测到网络恢复决策", 2)
    add_body(
        doc,
        """
        航空延误具有明显的系统性。一个机场的容量下降、前序航班晚到、局部天气冲击或航路拥堵，都会通过有向航线、飞机轮转、机组衔接和旅客中转关系传递到其他机场。若只回答“某一航班会不会延误”，仍无法处理恢复决策真正关心的问题：延误从哪里开始、沿哪些路径扩散、哪些节点值得优先投入资源、不同策略在成本和恢复效果之间如何取舍。

        因此，本文把研究对象界定为“面向突发扰动的机场网络运行恢复系统”。系统输入包括航班计划、机场/航线历史延误率、时间特征、航班量、外部冲击强度和策略预算；系统状态包括机场小时级延误水平、传播范围和最低性能；系统输出包括累计延误、恢复时间、传播范围、延误航班比例、成本和复杂度。这样的边界设置使预测模型、网络模型、仿真模型和决策模型可以在同一逻辑链中协同。
        """,
    )
    add_compact_matrix(
        doc,
        "系统概念与航空网络问题的对应",
        ["系统属性", "课程 PDF 含义", "本文中的具体体现"],
        [
            ["整体性", "系统整体功能不是各要素简单相加。", "单个航班是否延误并不等于全网恢复能力，必须看累计延误、传播范围和网络性能。"],
            ["关联性", "要素之间相互依赖、相互作用、相互制约。", "机场节点通过有向航线、前序晚到和流量转移形成传播关系。"],
            ["环境适应性", "系统功能会受外部环境变化影响。", "天气、高峰需求和枢纽容量下降被设置为外部情景扰动。"],
            ["功能性", "系统在环境作用中表现出特定功能。", "航空网络的功能是完成安全、准点、连通和可恢复的运输服务。"],
            ["层次性", "不同层次具有不同整体性。", "航班、机场、航线、机场群和全网构成多层次恢复问题。"],
            ["多元/多维性", "系统由不同性质部分和多种属性组成。", "旅客、航司、机场、监管者的目标不同，评价指标也同时覆盖效率、韧性、成本和复杂度。"],
        ],
        [1400, 3300, 4300],
        body_size=7.1,
        header_size=7.4,
    )
    add_heading(doc, "1.2 多主体目标冲突", 2)
    add_bullet(
        doc,
        [
            "旅客侧关注准点性、连接航班可靠性和信息透明度，倾向于缩短极端延误尾部。",
            "航空公司侧关注飞机/机组利用率、恢复成本和航班取消风险，倾向于控制资源投入。",
            "机场侧关注容量恢复、跑道/登机口瓶颈和枢纽连通性，倾向于保护关键节点。",
            "监管与公共管理侧关注安全、公平和全网稳定性，倾向于降低跨机场传播风险。",
        ],
    )
    add_formula(doc, "目标：min {累计延误, 恢复时间, 传播范围, 成本, 复杂度}; max {网络性能, 策略可解释性}", "系统工程视角强调多目标权衡；本项目不把单一指标作为唯一最优标准。")
    add_table_from_rows(
        doc,
        "系统边界、变量与反馈口径",
        ["层次", "变量/对象", "报告证据", "决策含义"],
        [
            ["外部环境", "天气扰动、高峰需求、枢纽容量下降、空域与机场运行约束。", "仿真情景 normal、peak、weather、hub_capacity；扰动项 w(t)。", "外部状态不可完全控制，只能通过情景和韧性策略降低影响。"],
            ["系统组成", "机场节点、航线边、承运航司、航班计划、历史运行状态。", "30 个机场、824 条有向边、机场/航线历史延误率。", "恢复策略必须兼顾节点位置和边的传播作用。"],
            ["系统状态", "机场小时级延误分钟、传播范围、最低性能、风险概率。", "传播矩阵 A、恢复曲线、热力图、风险预测概率 p_i。", "状态变量将预测模型与动态仿真连接起来。"],
            ["控制变量", "恢复预算、关键枢纽优先、统一缓冲、动态组合。", "四类策略在多情景下的累计延误、恢复时间、成本和复杂度。", "控制变量体现可行性约束下的方案综合。"],
            ["反馈输出", "策略排名、风险/不确定决策结果、Web Demo 可视化。", "TOPSIS 得分、期望损失、Hurwicz、后悔值、敏感性分析。", "通过反馈识别推荐策略的适用边界，而不是给出绝对最优。"],
        ],
        [1300, 2800, 2700, 2600],
    )
    add_labeled_points(
        doc,
        "本章与图片/表格的对应解释",
        [
            ("表格作用", "“系统边界、变量与反馈口径”不是装饰性表格，而是后文所有图表的索引：描述性图表对应外部环境和系统状态，网络图对应系统组成，仿真图对应控制变量，评价表对应反馈输出。"),
            ("系统边界", "报告只研究公开 BTS 数据能够支撑的机场网络与航班计划层级，不把机组、机尾号、实时天气雷达等不可得变量强行写入模型；这些内容在局限部分作为后续改进。"),
            ("管理含义", "一旦系统边界明确，后续策略推荐就不能被理解为真实航空公司可直接执行的排班指令，而是基于公开数据和系统工程方法得到的恢复策略比较。"),
        ],
    )

    add_heading(doc, "2 系统工程方法论映射", 1)
    add_body(
        doc,
        """
        课程 PDF/参考材料强调，系统工程方法论不是单一算法，而是一套面向复杂系统问题的思维方式、理论基础和程序。本文据此将“问题定义 - 模型化 - 优化 - 评价 - 决策 - 反馈”贯穿到项目流程中。霍尔三维结构提供总体组织框架；切克兰德方法论帮助识别多主体软问题；系统模型化与解释结构模型化技术(ISM)用于把复杂因素结构化；状态空间模型用于动态传播仿真；系统评价、AHP（层次分析法）、模糊综合评价和风险/不确定决策用于把方案差异转化为可解释选择。
        """,
    )
    add_table_from_rows(
        doc,
        "表 1  系统工程方法与本项目实现对应关系",
        ["方法", "课程要点", "本项目落地"],
        [
            ["霍尔三维结构", "从时间维、逻辑维、知识维组织系统工程任务。", "时间维对应数据准备、模型训练、仿真评估、交付部署；逻辑维对应问题定义、模型建立、综合评价、决策；知识维对应统计学习、复杂网络、控制/仿真和评价决策。"],
            ["切克兰德软系统方法", "用于目标模糊、主体诉求冲突的复杂管理问题。", "把旅客、航空公司、机场、监管者的目标差异写入系统边界，避免把“最低延误”简单等同于唯一目标。"],
            ["系统模型化", "用抽象模型反映主要要素、相互作用和因果关系，降低试验成本。", "用风险预测、有向网络、状态空间方程和策略评价矩阵替代不可直接试验的真实航空网络扰动。"],
            ["ISM 解释结构模型", "用有向图和矩阵处理要素关系，形成层级结构。", "将天气、容量、计划、前序晚到、拥堵、跨机场传播、恢复时间等因素整理为递阶关系。"],
            ["系统评价/AHP（层次分析法）/模糊综合评价", "评价多方案价值，为决策提供依据；AHP 将复杂问题分解并两两比较。", "构建 10 个评价指标，融合 AHP（层次分析法）与熵权，计算 TOPSIS；再用模糊等级解释策略质量。"],
            ["风险型与不确定型决策", "概率已知时用期望损失；概率未知时可用乐观准则、悲观准则、后悔值法(Savage准则/最小最大后悔值准则)、Hurwicz 折中准则等。", "对多情景损失矩阵分别计算期望损失、Laplace、Hurwicz、最小最大后悔值准则，明确推荐边界。"],
        ],
        [1600, 3300, 4100],
    )
    add_table_from_rows(
        doc,
        "系统工程五项原则在项目中的体现",
        ["原则", "课程含义", "报告实现", "对应证据"],
        [
            ["整体性", "由整体功能和总目标决定局部目标，避免局部最优损害系统效果。", "不只优化单航班预测准确率，而是把预测、网络、传播、评价和决策组合成闭环。", "摘要、图 12、表 5-8。"],
            ["有序相关性", "关注系统内部要素、子系统之间的相互关系。", "用机场有向网络、航线边权、ISM 邻接/可达矩阵描述要素关系。", "图 12、图 17-19。"],
            ["动态性", "注意系统内外变化的方向、速度和程度。", "用状态空间模型递推机场小时延误状态，并在多情景下观察恢复曲线。", "图 20-24。"],
            ["目标优化", "以最大效益、最小代价为出发点制定方案。", "评价累计延误、恢复时间、传播范围、成本、复杂度等多目标。", "表 5-6、图 25-28。"],
            ["可行性", "在可行方案中寻求满意或非劣方案。", "把预算、复杂度和风险态度纳入决策，保留动态组合与基准策略的条件边界。", "表 7-8、图 29-30。"],
        ],
        [1300, 2700, 3100, 1900],
    )
    add_compact_matrix(
        doc,
        "系统分析要素与本文问题结构",
        ["系统分析要素", "课程 PDF 表述", "本文落点"],
        [
            ["问题", "现实系统与目标系统之间存在偏差。", "真实航空网络在突发扰动下会出现延误扩散，目标系统应快速恢复且成本可控。"],
            ["目的/目标", "目的具有整体性，目标具有从属性和多样性。", "总体目的为提升网络韧性；分目标包括降低累计延误、缩短恢复时间、控制成本和复杂度。"],
            ["方案", "为达到目标提出若干预备方案。", "设置基准策略、统一缓冲、关键枢纽优先、动态组合四类方案。"],
            ["模型", "研究与解决问题的基本框架。", "使用预测模型、机场网络、ISM、状态空间模型和评价矩阵表达现实系统。"],
            ["评价", "比较不同方案对系统目的的达到程度。", "采用 AHP/熵权/TOPSIS、模糊评价、风险型和不确定型决策进行检验。"],
            ["决策者", "与系统分析人员有机配合。", "设定为兼顾旅客、航司、机场和监管视角的综合决策者。"],
        ],
        [1500, 3100, 4400],
    )
    add_body(
        doc,
        """
        这里特别需要说明，课程 PDF 中的原则不是被简单罗列，而是直接改变了本项目的建模口径。整体性原则要求我们把“预测一个航班是否延误”上升为“解释网络如何恢复”；有序相关性原则要求将机场和航线作为相互作用系统；动态性原则要求引入状态递推而不是只做静态排名；目标优化原则要求在收益和成本之间权衡；可行性原则则要求保留基准策略、统一缓冲、枢纽优先和动态组合等现实可解释方案，而不是只追求数学上最激进的单一方案。
        """,
    )
    add_heading(doc, "2.1 霍尔三维结构如何组织本文", 2)
    add_body(
        doc,
        """
        霍尔三维结构强调，系统工程任务可以同时从时间维、逻辑维和知识维展开。本文的时间维不是简单的写作顺序，而是从原始数据获取到静态 Demo 交付的工程流水线；逻辑维则体现为“识别问题 - 建立模型 - 求解比较 - 评价决策 - 反馈验证”；知识维把机器学习、复杂网络、运筹优化、控制仿真和综合评价放在同一报告中。这样处理的好处是，每一张图都能回答它在系统工程闭环中的位置，而不是成为孤立可视化。

        具体到本文，时间维从数据准备、模型训练、网络建模、传播仿真、评价排序到报告/Web 交付逐步推进；逻辑维对应课程 PDF 中的摆明问题、系统设计、方案综合、模型化、优化、决策和实施反馈；知识维则说明为什么需要同时使用统计学习、图论、状态空间、AHP（层次分析法）、模糊评价和决策准则。也就是说，本文不是先画图再找方法，而是先用三维结构规定每类图表承担的分析任务。
        """,
    )
    add_compact_matrix(
        doc,
        "霍尔三维结构在本文中的具体展开",
        ["维度", "课程含义", "本文落点", "对应证据"],
        [
            ["时间维", "系统工程工作阶段或进程。", "数据获取、清洗、训练、验证、仿真、评价、交付。", "数据审计表、模型指标、Web 截图。"],
            ["逻辑维", "摆明问题、系统设计、方案综合、模型化、优化、决策、实施。", "从延误传播问题界定到动态组合/基准策略的条件推荐。", "第 1-11 章主线。"],
            ["知识维", "完成任务所需的专业知识和方法。", "统计学习、复杂网络、ISM、状态空间、AHP/熵权/TOPSIS、风险决策。", "图 1、图 12-30、表 5-8。"],
        ],
        [1200, 3000, 3300, 1700],
    )
    add_heading(doc, "2.2 切克兰德方法与软约束", 2)
    add_body(
        doc,
        """
        航空恢复策略不是纯粹技术最优问题。若只按最低累计延误选择统一缓冲，可能造成过高成本；若只按成本选择基准策略，又会放大旅客体验和网络韧性风险。切克兰德软系统方法的价值在于提醒我们先识别不同主体的“世界观”和评价准则，再将冲突显式写入指标体系。本文把运行效率、网络韧性、航班影响、成本与复杂度共同纳入评价，就是对软系统问题的工程化表达。

        切克兰德方法论强调从问题及环境识别、根底定义、建立概念模型、与现实比较、寻求改善途径到实施评估的学习过程。本文将“理想概念模型”定义为一个能够在扰动中保持全网稳定、快速恢复且成本可控的机场网络；再用公开数据、预测误差、成本约束和多主体目标冲突与该理想模型比较。由此得到的不是单一“最优答案”，而是恢复优先、成本保守、风险规避等不同立场下的条件式策略建议。
        """,
    )
    add_compact_matrix(
        doc,
        "切克兰德软系统方法在本文中的应用",
        ["步骤", "本文处理", "产出"],
        [
            ["问题及环境识别", "识别天气、高峰需求、枢纽容量下降、航线连接和多主体目标冲突。", "系统边界与目标冲突表。"],
            ["根底定义", "把项目定义为“公开数据支撑下的机场网络恢复策略比较系统”。", "第 1 章研究对象与系统输入输出。"],
            ["概念模型", "建立风险预测、网络传播、状态仿真和多准则评价的理想链路。", "第 2 章方法链与图 1。"],
            ["现实比较", "用 BTS 数据缺口、模型 AUC、成本假设和情景有限性检验概念模型边界。", "第 3、5、11 章。"],
            ["改善与反馈", "给出动态组合/基准策略的条件推荐，并通过 Web Demo 和复现脚本回流检查。", "第 9-10 章与附录。"],
        ],
        [1800, 4800, 2400],
    )
    add_heading(doc, "2.3 方法链如何落到图表与结论", 2)
    add_body(
        doc,
        """
        本文的系统工程方法链可以理解为四个接口。第一个接口是“现实系统到数据系统”，即把真实航班运行抽象为机场、航线、时间、风险和延误状态。第二个接口是“数据系统到结构模型”，即通过有向网络和 ISM 把要素关系表达出来。第三个接口是“结构模型到动态模型”，即用状态空间方程进行扰动思想试验。第四个接口是“动态结果到决策”，即用系统评价和决策准则把多情景结果转化为可执行建议。这样的写法能保证每张图片都对应一个分析任务：描述图发现异质性，网络图解释传播位置，矩阵图解释结构关系，仿真图评估恢复路径，评价图支撑策略选择。
        """,
    )
    add_full_picture(
        doc,
        "fig_25_ism_hierarchy.png",
        "图 1  延误因素多级递阶结构",
        "该图是本文最核心的系统工程桥梁：底层根源因素包括天气、容量、计划和缓冲，中间层对应前序晚到、拥堵与航线连接，上层对应传播范围和恢复时间。它说明后续预测、网络和仿真模型都服务于同一个结构解释，而不是彼此独立。",
        width=6.0,
    )
    add_labeled_points(
        doc,
        "图 1 的阅读方法",
        [
            ("从下往上看", "底层因素更接近外部环境和计划约束，例如天气、容量和需求高峰；这些因素通常不能被恢复策略完全消除，只能被监测、缓冲或削弱。"),
            ("从中间层看", "前序晚到、地面拥堵和航线连接把局部问题转化为跨机场问题，这正是后文机场网络和传播矩阵需要存在的原因。"),
            ("从上层看", "传播范围、恢复时间和累计延误是管理者真正关心的系统结果，因此它们被纳入评价指标和 TOPSIS 决策矩阵。"),
        ],
    )

    add_heading(doc, "3 数据来源、审计与防泄漏设计", 1)
    add_body(
        doc,
        f"""
        本项目采用 BTS Airline On-Time Performance 数据，覆盖 {audit['date_min']} 至 {audit['date_max']}。原始记录 {audit['raw_rows']:,} 条，清洗后聚焦前 {kpi['airport_count']} 个机场的 {audit['scoped_rows_top_airports']:,} 条记录。取消率为 {audit['cancelled_rate']:.2%}，备降率为 {audit['diverted_rate']:.2%}。由于延误原因、实际到达时间、实际飞行时间等字段在计划阶段不可获得，模型训练时必须严格排除这些事后变量。

        防泄漏处理是本项目可信度的关键。训练、验证、测试按时间顺序划分，训练集 {audit['split_counts']['train']:,} 条、验证集 {audit['split_counts']['valid']:,} 条、测试集 {audit['split_counts']['test']:,} 条；航线、机场、航空公司历史延误率只由训练窗口估计并映射到验证/测试窗口。这样会降低表面指标，却更符合真实调度环境：决策者只能使用起飞前可知信息，而不能使用延误发生后的解释变量。

        从系统分析程序看，本章对应“认识问题”和“环境分析”。课程材料要求在认识问题阶段区分环境要素、划定系统边界；这里的边界就是 2024Q1 BTS 公开航班数据、样本内前 30 个机场、计划阶段可获得字段和可复现清洗规则。后续所有预测、网络、仿真和评价结果都必须回到这个边界解释，不能把缺少的实时天气、机尾号、机组轮转等变量假装成已经观测到的系统状态。
        """,
    )
    add_table_from_rows(
        doc,
        "表 2  数据资产与复现文件节选",
        ["文件", "记录数", "字段数", "说明"],
        [
            [row.file_name, row.rows, row.columns, "原始/中间数据资产，可追溯到数据清单。"]
            for row in manifest.head(6).itertuples(index=False)
        ],
        [3000, 1700, 1300, 3000],
    )
    add_table_from_rows(
        doc,
        "计划阶段建模字段处理规则",
        ["字段类别", "保留/剔除", "处理方式", "原因"],
        [
            ["计划字段", "保留", "计划起飞/到达时间、星期、月份、航线、机场、航司等进入特征。", "这些信息在航班运行前可知，符合部署条件。"],
            ["训练期历史风险", "保留", "仅用训练窗口计算航线、机场、机场小时、航司历史延误率。", "给模型提供经验风险，同时避免从验证/测试期偷看未来。"],
            ["实际运行字段", "剔除", "实际起飞、实际到达、实际飞行时间、到达延误分钟数不进入预测特征。", "这些变量在计划阶段不可知，使用会造成严重信息泄漏。"],
            ["延误原因字段", "剔除", "CarrierDelay、WeatherDelay、NASDelay 等只用于解释，不用于计划阶段预测。", "延误原因通常在事后归因，不能作为起飞前决策依据。"],
            ["取消/备降记录", "审计后处理", "统计取消率和备降率，主预测样本聚焦可评价到达延误的航班。", "保持目标定义清晰，避免取消/备降与到达延误混合解释。"],
        ],
        [1500, 1200, 4000, 2800],
    )
    add_labeled_points(
        doc,
        "数据章节的系统工程解释",
        [
            ("认识问题", "数据审计对应系统分析的起点：先确认研究对象、样本范围、可观测变量和不可观测变量，再决定能回答哪些问题。"),
            ("环境分析", "BTS 数据只覆盖公开航班运行记录，不能直接代表天气雷达、空域流控、机组排班和飞机轮转；这些缺口决定了模型必须保留边界说明。"),
            ("模型约束", "时间切分和防泄漏处理让模型更接近真实调度场景：未来不能被训练过程提前看到，事后变量不能伪装成计划阶段信息。"),
            ("评价基础", "若数据边界不清，后续 TOPSIS、风险决策和 Web 展示都会失去依据；因此本章先把数据资产、字段处理和复现路径写清楚。"),
        ],
    )
    add_picture_grid(
        doc,
        [
            ("fig_13_split_delay_rate.png", "图 2  时间切分后的目标比例", "测试期延误率与训练期不同，提示随机切分会高估泛化表现。"),
            ("fig_06_delay_distribution.png", "图 3  到达延误分钟数分布", "延误分钟数长尾明显，恢复策略应关注尾部严重延误压缩。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 2-3 对数据可信度的说明",
        [
            ("时间切分", "图 2 说明训练、验证、测试窗口的延误比例并不完全一致，因此必须用时间外测试衡量泛化，不能把未来窗口随机混入训练集。"),
            ("长尾分布", "图 3 说明大多数航班延误较短，但少数严重延误会显著拉高累计延误；这也是后文同时评价平均延误、累计延误和恢复损失面积的原因。"),
            ("审计价值", "数据审计不是形式环节，而是为了证明模型输入、目标变量和评价窗口均可复现，并为后续图表建立可信来源。"),
        ],
    )

    add_heading(doc, "4 描述性分析：延误的时间异质性与空间异质性", 1)
    add_body(
        doc,
        """
        描述性分析的目的不是“画图凑材料”，而是为系统建模找证据。若延误在时间、机场和航线之间近似均匀，复杂网络和状态传播模型的必要性就会降低；但图组显示，延误在日期、星期、小时、机场和航线维度均存在显著差异，这说明恢复资源不能平均撒布，而要结合时间窗口与节点位置进行分配。
        """,
    )
    add_picture_grid(
        doc,
        [
            ("fig_07_daily_volume_delay.png", "图 4  每日航班量与延误率趋势", "航班量与延误率并非完全同步，说明容量、天气、计划缓冲和网络传导共同作用。"),
            ("fig_08_week_hour_heatmap.png", "图 5  星期与小时延误率热力图", "高风险时段集中出现，是加入计划小时、星期和历史拥堵特征的依据。"),
            ("fig_09_volume_delay_scatter.png", "图 6  航班量与延误率关系", "高流量机场不必然最高延误，不能只按吞吐量排恢复优先级。"),
            ("fig_12_route_risk_bubble.png", "图 7  航线风险气泡图", "航线维度揭示了局部高风险连接，为边权设置和传播矩阵掩码提供依据。"),
        ]
    )
    add_callout(doc, "读图结论", "探索性分析支持三个建模判断：第一，预测模型需要时间和历史风险特征；第二，网络模型需要同时考虑节点与边；第三，策略评价需要跨情景比较，而不能只看平均延误。")
    add_body(
        doc,
        """
        对这些描述性图表的阅读要避免两个误区。第一，不能看到某个机场延误率高就直接判定它是恢复资源第一优先级，因为它可能只是局部高风险节点，对全网传播的中心性并不高。第二，不能看到某些时段延误率高就简单认为增加统一缓冲即可解决，因为缓冲会增加成本，也可能把有限资源消耗在低传播影响的节点上。系统工程强调的是把局部现象放回整体结构中解释：时间图说明何时风险升高，空间图说明在哪里发生，航线图说明如何扩散，后续模型才说明应该如何控制。
        """,
    )
    add_table_from_rows(
        doc,
        "描述性图组对后续模型的作用",
        ["图组", "揭示的问题", "后续模型使用"],
        [
            ["延误分布", "延误分钟数呈长尾，少数极端延误会显著影响均值。", "评价指标加入累计延误、最低性能和传播范围，避免只看平均表现。"],
            ["每日趋势", "航班量和延误率不同步，说明需求强度不是唯一解释。", "仿真情景同时设置高峰需求和天气/容量冲击。"],
            ["星期-小时热力图", "时段风险存在集中区域，运行状态具有时间异质性。", "预测模型加入计划小时、星期和机场小时历史风险。"],
            ["航线气泡图", "部分航线同时具有高频、高风险和高延误暴露。", "网络边权和传播掩码考虑航线级联系，而不是只看节点。"],
        ],
        [1600, 3600, 3800],
    )
    add_labeled_points(
        doc,
        "图 4-7 的综合解释",
        [
            ("时间异质性", "每日趋势和星期-小时热力图共同说明，延误风险不仅受机场影响，也受时间窗口影响；因此计划小时、星期和滚动历史状态进入模型是必要的。"),
            ("空间异质性", "机场散点和航线气泡图说明，机场流量、延误率和风险并非同一个概念；高流量机场可能更重要，高延误率机场也可能只是局部问题。"),
            ("建模过渡", "这些图把读者从“描述现象”带到“需要模型”：先看见异质性，再用预测模型量化风险，用网络模型表达联系，用仿真模型观察扩散。"),
        ],
    )

    add_heading(doc, "5 计划阶段延误风险预测", 1)
    add_body(
        doc,
        f"""
        预测目标为 ArrDel15，即到达延误是否达到 15 分钟。本文比较逻辑回归、随机森林和 LightGBM，采用时间外测试指标进行选择。若只看准确率，模型可能通过偏向多数类获得虚高结果；因此同时关注 ROC-AUC、PR-AUC、召回率、F1、Brier 分数和校准曲线。当前最佳模型为 {model['best_model']}，阈值 {model['threshold']:.2f}，测试 ROC-AUC {model['best_test_metrics']['roc_auc']:.3f}、PR-AUC {model['best_test_metrics']['pr_auc']:.3f}、Brier {model['best_test_metrics']['brier']:.3f}。

        从系统工程角度看，预测模型不是最终决策者，而是风险感知模块。其输出的概率 p_i 被输入机场节点画像、航线边权和恢复策略分配逻辑。即使单航班分类存在误报和漏报，只要风险排序能稳定地区分高风险时空片段，就能提升后续传播仿真和策略评价的解释力。
        """,
    )
    add_table_from_rows(
        doc,
        "表 3  模型时间外测试表现",
        ["模型", "阈值", "AUC", "PR-AUC", "召回率", "F1", "Brier"],
        [
            [r.model, r.threshold, r.test_roc_auc, r.test_pr_auc, r.test_recall, r.test_f1, r.test_brier]
            for r in model_metrics.itertuples(index=False)
        ],
        [1400, 1100, 1200, 1200, 1200, 1100, 1200],
    )
    add_formula(doc, "p_i = P(ArrDel15_i=1 | x_i)", "x_i 仅包含计划阶段可获得或由训练期统计得到的特征；事后延误原因、实际到达和实际飞行时间不进入主模型。")
    add_table_from_rows(
        doc,
        "预测模型输出在系统链路中的用途",
        ["输出", "解释", "后续用途"],
        [
            ["延误概率 p_i", "衡量某航班在计划阶段成为 ArrDel15 的相对风险。", "聚合为机场小时风险、航线风险和策略资源分配权重。"],
            ["排序能力 ROC/PR", "衡量模型能否把高风险样本排到前面，PR 更关注延误少数类。", "作为风险输入可信度依据，而不是单独作为最终管理目标。"],
            ["校准程度", "检验概率值是否接近真实频率，避免概率被系统性高估或低估。", "影响动态组合策略中的风险权重大小。"],
            ["特征重要性", "解释哪些计划变量、历史风险变量更影响风险排序。", "与 ISM 的底层因素和网络节点画像相互印证。"],
            ["混淆矩阵", "说明固定阈值下的误报和漏报结构。", "为实际部署中的阈值选择和人工复核提供参考。"],
        ],
        [1700, 3700, 3700],
    )
    add_picture_grid(
        doc,
        [
            ("fig_14_model_metrics.png", "图 8  模型指标对比", "随机森林在排序能力和综合表现上较均衡，适合作为后续风险输入。"),
            ("fig_16_roc_pr_curves.png", "图 9  ROC 与 PR 曲线", "延误样本不均衡时，PR 曲线更能反映高风险样本识别价值。"),
            ("fig_17_calibration_curve.png", "图 10  概率校准曲线", "校准用于检查概率是否可比较，避免策略优先级被失真概率放大。"),
            ("fig_18_shap_summary.png", "图 11  SHAP 特征重要性", "航线历史延误、机场小时风险和时间变量贡献较高，符合系统运行机理。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 8-11 的读图逻辑",
        [
            ("为什么不用准确率定胜负", "延误样本占比低于非延误样本，单看准确率会鼓励模型把更多航班判为正常；因此报告同时看 ROC-AUC、PR-AUC、召回率和 F1。"),
            ("为什么需要校准", "恢复策略使用的是概率排序和相对风险，不是单一分类标签；如果概率系统性偏高或偏低，资源分配强度会被扭曲。"),
            ("为什么看 SHAP", "SHAP 图将模型从黑箱输出拉回业务解释：航线历史延误、机场小时风险和时间因素较重要，说明模型学习到的是运行环境和网络状态，而非泄漏字段。"),
            ("对后续章节的输入", "预测概率在第 6 章被聚合为机场/航线风险，在第 8 章进入扰动状态，在第 9 章转化为评价指标的一部分。"),
        ],
    )
    add_compact_matrix(
        doc,
        "预测模型的误差如何进入系统决策",
        ["误差来源", "可能表现", "报告处理"],
        [
            ["不可观测外部因素", "天气、流控、机组、飞机轮转没有完整进入特征。", "在局限中说明，并把模型定位为风险排序器而不是自动调度器。"],
            ["类别不平衡", "非延误航班多，模型可能低估少数延误样本。", "同时报告 PR-AUC、召回率、F1 和混淆矩阵，不只看准确率。"],
            ["时间漂移", "测试期风险分布与训练期不同。", "采用时间外测试和历史窗口映射，避免随机切分导致的乐观偏差。"],
            ["概率偏差", "概率值可能高估或低估真实频率。", "使用校准曲线和 Brier 分数检查概率可用性。"],
            ["阈值选择", "不同阈值会改变误报/漏报平衡。", "将概率主要用于排序和聚合，固定阈值只作为解释辅助。"],
        ],
        [1800, 3600, 3600],
        body_size=7.1,
        header_size=7.4,
    )
    add_body(
        doc,
        """
        需要特别说明的是，AUC 中等并不意味着模型没有系统工程价值。单航班延误本身受天气、流控、机组和飞机周转等大量不可见因素影响，公开计划字段很难给出接近完美的分类结果。本文把模型定位为“风险排序器”而不是“自动调度器”：只要它能比随机排序更早识别高风险机场、航线和时段，就可以为状态传播模型和恢复策略提供更合理的初始条件。这个定位也符合模型化章节强调的原则：模型是现实系统的抽象和替代物，应抓住主要因素，同时承认简化边界。
        """,
    )

    add_heading(doc, "6 机场复杂网络与关键节点识别", 1)
    add_body(
        doc,
        f"""
        机场网络将机场视为节点，将 Origin-Dest 航线视为有向边。与只按航班量建边不同，本文的边权同时考虑航班量、平均延误和预测风险；节点关键性综合出港量、入港量、介数中心性、PageRank、平均风险和历史延误。网络摘要显示，样本网络包含 {network_summary['node_count']} 个节点、{network_summary['edge_count']} 条有向边，最大强连通分量覆盖全部节点，说明主要机场之间具有较强连通性，局部扰动具备跨区域扩散条件。

        综合关键性排名前五为 {'、'.join(network_summary['top_critical_airports'])}。这些机场之所以重要，不只是因为航班多，也因为它们在网络路径和风险叠加中具有更高传播影响。恢复策略若忽视中心性和传播位置，可能把资源投入到局部延误率高但全网影响较小的节点上。
        """,
    )
    top_nodes = nodes.sort_values("criticality", ascending=False).head(8)
    add_table_from_rows(
        doc,
        "表 4  关键机场指标节选",
        ["机场", "出港", "延误率", "平均风险", "介数", "PageRank", "关键性"],
        [
            [r.airport, r.departures, r.delay_rate, r.avg_risk, r.betweenness, r.pagerank, r.criticality]
            for r in top_nodes.itertuples(index=False)
        ],
        [900, 1300, 1200, 1300, 1200, 1300, 1500],
    )
    add_formula(doc, "K_i = omega_1 flow_i + omega_2 centrality_i + omega_3 risk_i + omega_4 delay_i", "K_i 表示机场综合关键性，用于把规模、网络位置、计划风险和历史延误归入同一排序口径。")
    add_table_from_rows(
        doc,
        "机场角色与恢复策略含义",
        ["角色", "识别依据", "策略含义"],
        [
            ["高中心-高风险节点", "中心性、流量、平均风险和历史延误均较高。", "优先配置恢复资源，防止局部冲击转化为全网传播。"],
            ["高中心-低风险节点", "网络通道作用强，但当前延误风险不一定最高。", "作为监控节点和备用恢复节点，重点观察冲击是否穿越。"],
            ["低中心-高风险节点", "局部延误率高，但对全网路径影响有限。", "采用局部缓冲和点状干预，避免过度占用全网资源。"],
            ["低中心-低风险节点", "规模、风险、中心性均不突出。", "维持常规监测，不作为主要恢复资源投放对象。"],
        ],
        [1800, 3600, 3600],
    )
    add_full_picture(
        doc,
        "fig_20_airport_network.png",
        "图 12  主要机场有向航线网络",
        "节点大小、颜色和边共同表达流量、风险与连接关系。该图说明恢复优先级不能只看单一机场指标，而要看节点在全网传播路径中的位置。",
        width=6.05,
    )
    add_picture_grid(
        doc,
        [
            ("fig_21_airport_criticality.png", "图 13  综合关键性排名", "MIA、DEN、DFW、FLL、MCO 位于前列，是恢复策略重点关注节点。"),
            ("fig_22_network_risk_scatter.png", "图 14  中心性-风险-航班量散点", "可区分“高风险但低中心”和“高中心但风险中等”的不同节点角色。"),
            ("fig_26_airport_role_quadrant.png", "图 15  机场角色象限", "象限图将节点划分为优先恢复、监控、防扩散和普通节点，便于策略解释。"),
            ("fig_10_airport_delay_rank.png", "图 16  机场延误率排名", "延误率排名与关键性排名不完全一致，说明网络位置必须纳入决策。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 12-16 对策略的含义",
        [
            ("网络图", "图 12 用整体拓扑说明样本机场不是孤立点，而是强连通的有向网络；强连通意味着局部恢复失败可能通过多条路径影响其他机场。"),
            ("关键性排名", "图 13 把流量、中心性、风险和历史延误合成统一口径，避免只按航班量或只按延误率做片面排序。"),
            ("角色象限", "图 14-15 将机场分成不同管理角色：高中心高风险节点优先投入，高风险低中心节点局部治理，高中心低风险节点加强监控。"),
            ("与图 16 的差异", "图 16 提醒我们，延误率高不等于系统关键性最高；这正是系统工程区别于单指标统计排名的地方。"),
        ],
    )
    add_body(
        doc,
        """
        网络章节对应课程 PDF 中“系统由要素、关系和环境构成”的观点。机场是要素，航线是关系，天气、高峰需求和容量变化是环境；只有把三者合在一起，才能解释为什么相同延误分钟数在不同机场具有不同系统后果。一个低中心节点发生延误，可能主要影响自身和少量航线；一个高中心节点发生延误，即使当前延误率不最高，也可能通过中转、航线密度和上游晚到扩大影响。因而关键性排名不是替代业务经验，而是把业务经验结构化、定量化，供后续策略分配使用。
        """,
    )

    add_heading(doc, "7 ISM 结构分析与模型化逻辑", 1)
    add_body(
        doc,
        """
        课程材料指出，ISM 的核心是用有向图、矩阵和可达关系处理复杂系统要素之间的结构关系，并形成递阶模型。本文把 ISM 用作“解释框架”而非强因果证明：它不声称某个因素必然导致另一个因素，而是帮助读者理解为什么天气、容量、计划缓冲、前序晚到和拥堵会逐层影响传播范围与恢复时间。

        在报告中，ISM 起到三重作用。第一，它把鱼骨图式因素识别转化为可计算的邻接矩阵和可达矩阵；第二，它解释为什么预测特征、网络指标和仿真状态变量是合理的；第三，它使后续综合评价不是单纯指标堆叠，而是来自系统结构的目标分解。

        具体建模时，因素集合可记为 S={天气冲击、需求高峰、机场容量、计划缓冲、前序晚到、地面拥堵、航线连接、传播范围、恢复时间、旅客影响}。邻接矩阵 A 表示因素之间是否存在直接影响，可达矩阵 M 表示经过若干传递路径后能否到达。课程 PDF 强调 ISM 的技术核心是通过可达矩阵处理建立递阶结构模型；本文据此把“根源因素 - 传导因素 - 状态因素 - 结果因素”分层，使状态空间模型和多准则评价具有可解释的结构来源。
        """,
    )
    add_picture_grid(
        doc,
        [
            ("fig_23_ism_adjacency.png", "图 17  ISM 邻接矩阵", "邻接矩阵表达因素之间是否存在直接影响，是结构模型的输入。"),
            ("fig_24_ism_reachability.png", "图 18  ISM 可达矩阵", "可达矩阵揭示直接与间接影响链，是递阶层级划分的基础。"),
        ]
    )
    add_formula(doc, "M = (A + I)^r, r <= n-1", "按照结构模型化思想，邻接矩阵 A 加单位矩阵 I 后进行布尔运算，可得到反映直接与间接影响链的可达矩阵 M。")
    add_compact_matrix(
        doc,
        "ISM 建模步骤与本文证据",
        ["步骤", "课程方法", "本文实现"],
        [
            ["识别要素", "提取问题构成要素。", "从延误机理、数据字段和运行逻辑中确定天气、容量、计划、前序晚到、拥堵、传播等因素。"],
            ["定义关系", "以二元关系描述要素间直接影响。", "构造 ISM 邻接矩阵，1 表示存在直接影响，0 表示未设定直接影响。"],
            ["计算可达", "用布尔运算得到直接与间接影响链。", "形成可达矩阵并据此判断因素层级。"],
            ["层级解释", "建立递阶结构模型并用文字说明。", "将因素划分为根源层、传导层、状态层和结果层，并连接到图 1、图 17-18。"],
            ["回到模型", "结构模型服务于后续分析。", "底层因素进入预测和情景，中层因素进入网络与传播矩阵，上层结果进入评价指标。"],
        ],
        [1400, 2600, 5000],
    )
    add_table_from_rows(
        doc,
        "ISM 因素层级的业务解释",
        ["层级", "代表因素", "解释作用"],
        [
            ["根源层", "天气冲击、机场容量、计划缓冲、需求高峰。", "这些因素多为外部或计划前约束，是后续延误状态变化的起点。"],
            ["传导层", "前序航班晚到、地面拥堵、航线连接、资源占用。", "这些因素把单点扰动转化为跨机场、跨时段传播。"],
            ["状态层", "机场小时延误水平、传播范围、最低性能。", "这些因素进入状态空间模型，描述系统在每一步的运行表现。"],
            ["结果层", "恢复时间、累计延误、旅客影响、策略成本。", "这些因素进入系统评价，最终影响恢复策略选择。"],
        ],
        [1400, 3400, 4200],
    )
    add_callout(
        doc,
        "ISM 与后续模型的关系",
        "底层因素进入预测特征和扰动情景，中层因素进入网络边权和传播矩阵，高层结果进入恢复时间、传播范围和旅客影响等评价指标。这样，图 17-18 不只是课程方法展示，而是模型链路的结构解释。",
    )
    add_labeled_points(
        doc,
        "图 17-18 与课程 ISM 方法对应",
        [
            ("邻接矩阵", "图 17 对应课程中“以二元关系表达要素之间是否存在直接影响”的步骤，矩阵中的 1 表示直接影响关系，0 表示未设定直接影响。"),
            ("可达矩阵", "图 18 对应对邻接矩阵加入单位矩阵后进行布尔运算的结果，反映要素之间的直接和间接影响链。"),
            ("递阶结构", "可达矩阵不是终点，它服务于因素层级划分：底层因素解释外部扰动来源，中层因素解释传播路径，上层因素解释评价结果。"),
            ("方法边界", "ISM 的作用是结构化认识问题，不替代统计因果识别；因此本文把它用于解释和建模组织，而非宣称所有边都具有严格因果强度。"),
        ],
    )
    add_compact_matrix(
        doc,
        "ISM 与预测、网络、仿真的变量承接",
        ["ISM 层级", "对应变量", "进入的模型环节", "解释价值"],
        [
            ["根源因素", "天气、容量、需求高峰、计划缓冲。", "情景扰动 w(t)、计划时间特征、缓冲策略。", "解释外部环境和计划约束为何是延误传播起点。"],
            ["传导因素", "前序晚到、航线连接、地面拥堵。", "机场网络边权、历史延误率、传播矩阵掩码。", "说明局部延误如何跨机场、跨时段传递。"],
            ["状态因素", "机场小时延误、传播范围、最低性能。", "状态向量 x(t)、恢复曲线、传播范围曲线。", "把静态结构转化为动态过程。"],
            ["结果因素", "累计延误、恢复时间、旅客影响、成本。", "TOPSIS 评价矩阵、风险/不确定决策损失矩阵。", "将仿真输出转化为可比较的策略选择。"],
        ],
        [1400, 2600, 2600, 2800],
        body_size=7.0,
        header_size=7.3,
    )

    add_heading(doc, "8 状态空间传播模型与扰动仿真", 1)
    add_body(
        doc,
        f"""
        状态传播模型以机场小时级正延误分钟数为状态变量。模型形式为 x(t+1)=A x(t)+B u(t)+G w(t)+epsilon(t)，其中 A 表示机场之间的延误传播，u(t) 表示恢复控制，w(t) 表示外部冲击。为避免小样本回归形成不现实的全连接传播，A 的估计使用网络掩码，只允许自身和存在航线连接的上游机场进入。原始谱半径为 {prop['spectral_radius_raw']:.3f}，经稳定化缩放后为 {prop['spectral_radius_final']:.3f}。

        仿真比较 normal、peak、weather、hub capacity 等情景，并在相同初始状态、相同冲击和相同预算口径下测试基准、统一缓冲、关键枢纽优先和动态组合。这样做可以避免把参数不一致误判为策略优劣，也符合系统工程中“模型求解 - 检验修订 - 综合评价”的流程。
        """,
    )
    add_formula(doc, "x(t+1)=A x(t)+B u(t)+G w(t)+epsilon(t)", "x(t) 为机场状态向量；A 为传播矩阵；u(t) 为恢复资源；w(t) 为天气、高峰和容量下降等扰动。")
    add_table_from_rows(
        doc,
        "状态空间传播模型变量解释",
        ["符号", "含义", "估计/设置方式", "管理解释"],
        [
            ["x(t)", "第 t 个小时各机场延误状态向量。", "由机场小时级正延误分钟聚合得到。", "表示系统当前运行压力。"],
            ["A", "机场之间的延误传播矩阵。", "带航线网络掩码估计，并通过谱半径稳定化。", "表示上游节点对下游节点的影响强度。"],
            ["u(t)", "恢复控制向量。", "由基准、统一缓冲、关键枢纽优先、动态组合四类策略生成。", "表示可控资源投入。"],
            ["w(t)", "外部扰动向量。", "由 normal、peak、weather、hub_capacity 等情景设定。", "表示决策者不可完全控制的自然状态。"],
            ["epsilon(t)", "残差与未建模因素。", "用模型误差吸收公开数据不可见因素。", "提醒结论应保留不确定性边界。"],
        ],
        [900, 2400, 3000, 2900],
    )
    add_body(
        doc,
        """
        状态空间模型的优势在于把“当前状态、控制变量和外部扰动”放在同一递推框架内。x(t) 不是单个航班的标签，而是机场网络的状态向量；A 不是普通相关系数表，而是经过网络掩码约束的传播关系；u(t) 不是事后解释，而是恢复策略可以施加的控制；w(t) 则对应天气、高峰和容量下降等自然状态。通过这种拆分，报告可以清楚区分“系统自身传播”“人为恢复控制”和“外部环境冲击”，也便于后续将仿真结果转换为系统评价指标。
        """,
    )
    add_compact_matrix(
        doc,
        "仿真情景与策略设置",
        ["类别", "名称", "含义", "评价关注点"],
        [
            ["情景", "normal", "常态运行扰动，用作基准压力测试。", "检验策略在普通波动下是否过度投入。"],
            ["情景", "peak", "高峰需求放大机场小时状态压力。", "检验资源配置是否能抑制拥堵扩散。"],
            ["情景", "weather", "天气扰动造成外部冲击项增强。", "检验策略对不可控自然状态的韧性。"],
            ["情景", "hub_capacity", "关键枢纽容量下降，冲击集中在高中心节点。", "检验网络关键节点保护能力。"],
            ["策略", "基准策略", "不额外施加强恢复控制。", "作为低成本参照。"],
            ["策略", "统一缓冲", "对全网统一增加缓冲。", "观察恢复收益与成本压力。"],
            ["策略", "关键枢纽优先", "优先保护高关键性节点。", "观察节点优先级是否降低传播。"],
            ["策略", "动态组合", "结合风险、关键性和状态变化动态分配资源。", "检验综合策略的恢复效率与复杂度。"],
        ],
        [900, 1400, 4000, 3100],
    )
    add_picture_grid(
        doc,
        [
            ("fig_27_propagation_matrix.png", "图 19  传播矩阵 A", "矩阵非零项主要集中在存在运行联系的机场对，避免无意义全连接。"),
            ("fig_28_recovery_heatmap.png", "图 20  策略-情景恢复热力图", "热力图展示不同策略在不同情景下的恢复差异，是决策矩阵来源。"),
            ("fig_29_recovery_curves.png", "图 21  恢复曲线", "动态组合恢复较快，但成本和复杂度更高，需进入多准则评价。"),
            ("fig_30_scenario_delay_compare.png", "图 22  多情景累计延误对比", "不同冲击下策略排序不完全一致，说明必须做情景稳健性检查。"),
            ("fig_31_spread_range_curves.png", "图 23  传播范围曲线", "传播范围用于衡量局部扰动是否扩散为全网问题。"),
            ("fig_32_cost_delay_pareto.png", "图 24  成本-延误 Pareto 关系", "统一缓冲效果好但成本高，基准成本低但延误大，动态组合位于折中区域。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 19-24 的仿真解释",
        [
            ("传播矩阵 A", "图 19 是状态空间模型的核心。矩阵越强，说明上游机场状态对下游机场下一时段延误影响越大；谱半径稳定化保证仿真不会在有限冲击下无限发散。"),
            ("策略热力图", "图 20 将策略与情景交叉展示，能看出同一策略在 normal、peak、weather、hub failure 等自然状态下表现不同。"),
            ("恢复曲线", "图 21 关注时间过程而非终点值。动态组合和统一缓冲恢复快，但成本、复杂度和实施难度不同，因此不能只凭曲线最低点选方案。"),
            ("情景对比", "图 22-23 说明天气和枢纽故障情景更容易产生大范围传播，恢复策略的稳健性必须跨情景验证。"),
            ("Pareto 取舍", "图 24 将延误收益和策略成本放在一起看，体现系统工程中的目标优化与可行性结合：低延误、高成本和低成本、高延误之间需要折中。"),
        ],
    )
    add_compact_matrix(
        doc,
        "仿真输出如何转化为评价指标",
        ["仿真输出", "计算含义", "进入评价的理由"],
        [
            ["累计延误", "整个仿真窗口内延误压力的总量。", "衡量系统总体损失，体现整体性目标。"],
            ["恢复时间", "状态恢复到阈值以内所需的时间。", "衡量策略对动态恢复过程的改善。"],
            ["最低性能", "冲击过程中系统性能下降的最低点。", "衡量网络韧性的抗冲击能力。"],
            ["传播范围", "受影响机场或时段的覆盖范围。", "衡量局部扰动是否扩散为全网问题。"],
            ["策略成本", "缓冲、资源优先和动态调配带来的相对投入。", "体现可行性和经济约束。"],
            ["复杂度", "策略解释、实施和协调难度。", "防止模型给出难以执行的理论最优方案。"],
        ],
        [1700, 3400, 3900],
        body_size=7.0,
        header_size=7.3,
    )
    add_body(
        doc,
        """
        状态空间仿真在本文中承担“思想试验”的作用。真实航空网络不能为了课程项目故意制造天气冲击或关闭枢纽容量，但模型可以在可控假设下比较不同策略的相对表现。课程 PDF 关于模型化的提示在这里尤其重要：模型应反映系统主要组成和相互作用，但模型也有局限，不能替代真实系统。本文因此把仿真结果转化为相对评价指标，而不是声称某个策略能在真实运营中精确减少多少分钟延误。
        """,
    )

    add_heading(doc, "9 多准则系统评价与策略决策", 1)
    add_body(
        doc,
        """
        系统评价的任务是对多种方案的价值进行定性或定量评定，为决策提供依据。本文将评价对象定义为四类恢复策略，评价主体设定为兼顾运行方、机场方和监管视角的综合决策者，评价目标包括运行效率、网络韧性、航班影响、成本与复杂度。由于这些指标既有收益型也有成本型，先进行同向化与归一化，再进行权重融合和 TOPSIS 排名。

        AHP（层次分析法）提供主观偏好，反映“恢复效率和网络韧性优先”的课程情景设定；熵权提供客观离散度，反映不同指标实际区分方案的能力。最终权重以 lambda=0.9 融合 AHP 与熵权，AHP 一致性检验 CR=0.000，说明判断矩阵内部一致。TOPSIS 接近度用于衡量方案距离正理想解与负理想解的相对位置。
        """,
    )
    add_table_from_rows(
        doc,
        "系统评价 5W1H 设置",
        ["要素", "本项目设定", "说明"],
        [
            ["What 评价对象", "基准策略、统一缓冲、关键枢纽优先、动态组合。", "四类策略代表从低成本到高响应的不同恢复方案。"],
            ["Who 评价主体", "兼顾机场运行、航司恢复、旅客体验和监管稳定性的综合决策者。", "对应切克兰德方法中多主体效用观点差异。"],
            ["Why 评价目的", "在突发扰动下选择兼顾效率、韧性、成本和可实施性的恢复策略。", "避免只按累计延误或只按成本单指标决策。"],
            ["When 评价时期", "基于 2024Q1 样本和多情景仿真结果进行期末/方案评价。", "可以作为后续实时系统的离线验证基线。"],
            ["Where 评价范围", "样本内前 30 个主要机场组成的有向航线网络。", "保证网络规模可解释、可复现、可视化。"],
            ["How 评价方法", "AHP（层次分析法）/熵权融合、TOPSIS、模糊综合评价、风险型与不确定型决策。", "既有主观偏好，也有客观离散度和决策态度检验。"],
        ],
        [1500, 4200, 3300],
    )
    add_compact_matrix(
        doc,
        "评价指标体系与准则层",
        ["准则层", "指标示例", "方向", "解释"],
        [
            ["运行效率", "累计延误、平均延误、恢复时间。", "越小越好", "衡量策略能否快速压缩延误。"],
            ["网络韧性", "最低性能、恢复损失面积、传播范围。", "性能越大越好；损失越小越好", "衡量冲击后系统保持和恢复功能的能力。"],
            ["航班影响", "延误航班比例、取消代理。", "越小越好", "将旅客体验和运行影响纳入评价。"],
            ["经济可行性", "相对成本。", "越小越好", "防止策略只追求效果而忽视资源约束。"],
            ["实施复杂度", "策略复杂度。", "越小越好", "反映方案能否在真实管理环境中解释和执行。"],
        ],
        [1500, 2900, 1800, 3100],
    )
    add_table_from_rows(
        doc,
        "表 5  主偏好下策略 TOPSIS 排名",
        ["策略", "得分", "累计延误", "恢复时间", "最低性能", "成本", "排名"],
        strategy_rows(rankings),
        [1400, 1100, 1400, 1300, 1300, 1200, 900],
    )
    weights_view = weights[["label", "criterion", "type", "ahp_weight", "entropy_weight", "combined_weight"]].head(10)
    add_table_from_rows(
        doc,
        "表 6  指标权重节选",
        ["指标", "准则层", "类型", "AHP", "熵权", "综合权重"],
        [list(row) for row in weights_view.itertuples(index=False, name=None)],
        [1600, 1300, 1000, 1200, 1200, 1400],
    )
    add_formula(doc, "C_i = D_i^- / (D_i^+ + D_i^-)", "TOPSIS 接近度 C_i 越大，表示方案越接近正理想解、远离负理想解。")
    add_formula(doc, "w = lambda * w_AHP + (1-lambda) * w_entropy, lambda=0.9", "AHP 表示恢复优先的主观偏好，熵权表示指标离散度；融合权重既保留课程情景偏好，也避免完全主观。")
    add_compact_matrix(
        doc,
        "AHP/熵权/TOPSIS 计算流程",
        ["步骤", "处理", "输出"],
        [
            ["指标同向化", "把成本型指标转换为收益方向或在距离计算中按成本型处理。", "可比较的评价矩阵。"],
            ["AHP 权重", "按目标层、准则层、指标层构造判断矩阵，并进行一致性检验。", "主观偏好权重。"],
            ["熵权修正", "根据指标在策略之间的离散度计算客观权重。", "指标区分度权重。"],
            ["权重融合", "以 lambda=0.9 融合 AHP 与熵权。", "综合权重表。"],
            ["TOPSIS 排序", "计算各方案到正理想解和负理想解的距离。", "接近度 C_i 与策略排名。"],
            ["模糊评价", "将连续得分映射到优秀、良好、一般、较差等等级。", "便于答辩解释的质量等级。"],
        ],
        [1400, 4900, 2700],
    )
    add_picture_grid(
        doc,
        [
            ("fig_33_strategy_radar.png", "图 25  策略雷达图", "雷达图显示动态组合在恢复收益上更均衡，但复杂度也较高。"),
            ("fig_34_topsis_score.png", "图 26  TOPSIS 综合得分", "动态组合排名第一，统一缓冲紧随其后，说明策略差异需要结合成本解释。"),
            ("fig_35_fuzzy_grades.png", "图 27  模糊综合评价等级", "动态组合和统一缓冲主要落在“良好”，基准策略主要落在“一般”。"),
            ("fig_36_indicator_weights.png", "图 28  指标权重分布", "运行效率与网络韧性权重较高，符合恢复优先的主偏好设定。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 25-28 的评价解释",
        [
            ("雷达图", "图 25 适合观察策略形态：动态组合在效率和韧性上更均衡，基准策略成本低但恢复表现弱，统一缓冲效果好但成本高。"),
            ("TOPSIS", "图 26 将多指标压缩为接近度得分，动态组合最高，说明在主偏好下它最接近正理想解、远离负理想解。"),
            ("模糊评价", "图 27 把连续指标转成“优秀、良好、一般、较差”等等级，更适合课堂答辩中解释方案质量，而不是只报小数。"),
            ("权重分布", "图 28 展示 AHP（层次分析法）/熵权融合后的评价偏好：运行效率和网络韧性权重较高，经济性和复杂度作为可行性约束参与折中。"),
        ],
    )
    add_compact_matrix(
        doc,
        "为什么需要多种评价方法交叉验证",
        ["方法", "优点", "局限", "本文采用方式"],
        [
            ["AHP（层次分析法）", "能表达课程情景和管理偏好，层次清楚。", "带有主观判断，需要一致性检验。", "用于给恢复效率和韧性更高权重，并报告 CR。"],
            ["熵权", "反映指标数据本身的差异度。", "不代表管理偏好，可能放大波动大的指标。", "作为客观权重修正，与 AHP 融合。"],
            ["TOPSIS", "能同时考虑正理想解和负理想解。", "结论依赖指标、权重和归一化口径。", "作为主偏好下的综合排序工具。"],
            ["模糊综合评价", "适合把连续得分解释为等级。", "等级边界有设定性，不宜替代数值排序。", "用于把策略质量解释为优秀、良好、一般、较差。"],
            ["风险/不确定决策", "能检验不同概率假设和风险态度。", "若损失矩阵偏重成本，推荐会更保守。", "作为 TOPSIS 的边界检验，不直接推翻主推荐。"],
        ],
        [1600, 2300, 2500, 3000],
        body_size=6.9,
        header_size=7.2,
    )
    add_body(
        doc,
        """
        从评价结果看，动态组合并非在所有单项指标上都绝对最好，而是在主偏好权重下实现了较好的综合平衡。统一缓冲的恢复效果接近，但相对策略成本明显更高；关键枢纽优先成本较低、方向明确，但对全网冲击的覆盖不如动态组合；基准策略几乎不增加成本，却在累计延误和恢复时间上付出较大代价。因此 TOPSIS 的价值在于把“好在哪里、代价是什么”放进同一评价空间，而不是简单替代专业判断。
        """,
    )
    add_heading(doc, "9.1 风险型与不确定型决策补充", 2)
    add_body(
        doc,
        """
        课程决策分析中，风险型决策适用于自然状态概率可以估计的情形；不确定型决策适用于概率不可知或决策者更关注态度准则的情形。本文将多情景仿真损失矩阵同时代入两类方法，目的不是推翻 TOPSIS，而是检验推荐方案在不同管理偏好下是否稳定。

        风险型决策使用期望值法：当 normal、peak、weather、hub_capacity 等自然状态可以给出概率时，比较每个策略的期望损失，若目标是费用或损失最小，则选择期望损失最小的方案。不确定型决策不假设概率，而是分别代入乐观法、悲观法、Laplace 等可能性、Savage 后悔值法和 Hurwicz 折中准则。这样可以回答一个非常实际的问题：如果评委追问“概率不知道怎么办”或“如果决策者更保守怎么办”，报告仍有方法依据，而不是只依赖主偏好 TOPSIS 排名。
        """,
    )
    add_formula(doc, "E(L_s)=sum_c P(c) * L(s,c)", "风险型决策期望损失：策略 s 在情景 c 下的损失为 L(s,c)，自然状态概率为 P(c)。")
    add_formula(doc, "R(s,c)=L(s,c)-min_s L(s,c),   choose min_s max_c R(s,c)", "Savage 后悔值法：先计算每个情景下相对最优方案的后悔值，再选择最大后悔值最小的策略。")
    add_compact_matrix(
        doc,
        "决策准则含义与适用边界",
        ["准则", "课程含义", "本文解释"],
        [
            ["期望损失", "概率可估计时比较各方案期望益损值。", "情景概率来自历史经验或专家判断时，选择期望损失较低方案。"],
            ["Laplace 等概率", "概率未知时假设各自然状态等可能。", "作为中性基线，检验策略是否依赖强概率假设。"],
            ["Hurwicz 折中", "用乐观系数在最好和最坏结果间折中。", "反映风险态度；乐观系数变化会改变推荐边界。"],
            ["Savage 后悔值", "先求后悔矩阵，再选择最大后悔值最小方案。", "适用于难以承受选错策略代价的保守运营情景。"],
        ],
        [1800, 3500, 3700],
        body_size=7.1,
        header_size=7.4,
    )
    add_compact_matrix(
        doc,
        "TOPSIS 与决策分析的关系",
        ["方法", "回答的问题", "本文结论含义"],
        [
            ["TOPSIS", "在既定权重和多指标评价矩阵下，哪个方案最接近正理想解。", "恢复效率和网络韧性优先时，动态组合综合表现最好。"],
            ["风险型期望损失", "情景概率可估计时，哪个方案的期望损失更低。", "成本惩罚较强时，基准策略更稳妥。"],
            ["Laplace", "概率未知但各自然状态等可能时，哪个方案平均损失较低。", "检验结论是否过度依赖主观概率。"],
            ["Hurwicz", "乐观和悲观之间折中时，推荐是否变化。", "反映管理者风险态度对方案选择的影响。"],
            ["Savage 后悔值", "事后看错选方案的最大遗憾能否被控制。", "适用于不愿承担错误投入成本的保守运营方。"],
        ],
        [1700, 3900, 3500],
    )
    add_table_from_rows(
        doc,
        "表 7  风险型决策期望损失",
        ["策略", "期望损失", "风险排名"],
        [[format_strategy(r.strategy), r.expected_loss, int(r.risk_rank)] for r in risk.itertuples(index=False)],
        [2500, 2800, 1800],
    )
    add_table_from_rows(
        doc,
        "表 8  不确定型决策准则结果",
        ["准则", "推荐策略", "准则值"],
        [[r.criterion, format_strategy(r.recommended_strategy), r.value] for r in uncertainty.itertuples(index=False)],
        [3200, 2300, 2500],
    )
    add_picture_grid(
        doc,
        [
            ("fig_37_weight_sensitivity.png", "图 29  权重敏感性分析", "恢复偏好下降时基准策略可能反超，说明结论存在明确适用边界。"),
            ("fig_38_risk_vs_topsis.png", "图 30  风险决策与 TOPSIS 对照", "TOPSIS 偏综合效益，风险型决策偏保守成本，两者共同构成推荐边界。"),
        ]
    )
    add_labeled_points(
        doc,
        "图 29-30 的决策解释",
        [
            ("权重敏感性", "图 29 检验“如果决策者没有这么重视恢复效率，推荐是否还成立”。当恢复权重下降、成本权重上升时，基准策略可能反超，说明推荐有条件。"),
            ("风险型决策", "图 30 对比期望损失和 TOPSIS。TOPSIS 偏向综合效益，风险型决策在损失矩阵下更保守，体现不同决策准则会给出不同方案。"),
            ("不确定型决策", "乐观、悲观、Laplace、Hurwicz 和最小最大后悔值准则在本项目中都指向基准策略，原因是损失矩阵把成本惩罚纳入得更重。"),
            ("最终口径", "因此报告不写“动态组合永远最优”，而写“恢复优先时推荐动态组合；成本极端保守或概率不可知时保留基准策略”。"),
        ],
    )
    add_compact_matrix(
        doc,
        "从决策准则到推荐口径的转换",
        ["决策态度", "采用准则", "可能推荐", "报告表述"],
        [
            ["恢复优先", "AHP/熵权/TOPSIS 主偏好。", "动态组合。", "适合强调网络韧性、恢复时间和旅客影响的场景。"],
            ["成本保守", "风险型期望损失、悲观或后悔值准则。", "基准策略或低强度策略。", "适合预算紧、无法承担额外资源投入的场景。"],
            ["概率未知", "Laplace、Hurwicz、Savage 等不确定型准则。", "视乐观系数和损失矩阵而定。", "需要明确自然状态概率未知时的管理态度。"],
            ["现场答辩", "主推荐 + 边界说明。", "条件式推荐。", "先说明动态组合为什么好，再说明何时不应盲目采用。"],
        ],
        [1500, 2500, 1900, 3500],
        body_size=7.0,
        header_size=7.3,
    )
    add_para(
        doc,
        f"最终推荐口径：当目标偏向恢复效率、网络韧性和旅客影响时，推荐 {format_strategy(decision['recommended_strategy'])}；当预算极紧、决策者按期望损失、悲观准则或最小最大后悔值准则保守选择时，{format_strategy(decision['risk_recommended_strategy'])} 更稳妥。报告采用有条件推荐，避免把多目标决策写成单目标最优。",
        size=9.1,
        color=NAVY,
        bold=True,
        before=2,
        after=4,
        line=1.14,
    )

    add_heading(doc, "10 静态 Web Demo 与工程化交付", 1)
    add_body(
        doc,
        """
        为了让报告结果能够被现场追问和复核，项目将核心表、矩阵和图形导出为静态 Web Demo。静态站不依赖 Python 后端，数据集中在 `web/assets/data/flightresilience-data.json`，页面按“总览 - 数据 - 预测 - 网络 - 仿真 - 决策 - 方法”组织。新版页面删除了单独现场入口区，把首屏、方法区和尾页整合为演示型控制台；自动演示不再只是向下滚动，而是按讲稿真实触发机场选择、航线气泡、滑杆、网络节点、情景策略选项和风险-TOPSIS 点。

        这一设计使评委可以从任意图表回到数据来源和方法链，而不是只观看不可交互的截图。对于答辩者而言，自动演示承担“标准讲述路径”，手动交互承担“追问复核路径”：先用 25 秒左右走完整链路，再根据评委问题进入数据、预测、网络、仿真或决策页面。
        """,
    )
    add_table_from_rows(
        doc,
        "Web Demo 页面与报告证据关系",
        ["页面", "展示内容", "支撑的报告章节"],
        [
            ["总览", "样本范围、核心 KPI、主推荐策略、讲解进度条和标准演示入口。", "摘要、结论、策略推荐。"],
            ["数据/预测", "数据审计、时间切分、模型指标、ROC/PR、校准、SHAP，以及航线气泡和滑杆风险估算。", "第 3-5 章。"],
            ["网络", "机场有向网络、Top 关键性条形、机场画像和传播关系；点击节点会联动画像。", "第 6-7 章。"],
            ["仿真", "情景/策略快捷选项、恢复热力图、恢复曲线、传播范围、成本-延误权衡。", "第 8 章。"],
            ["决策/方法", "TOPSIS、权重敏感性、风险准则、完整公式块、系统工程方法链和尾页总结。", "第 9 章及方法论部分。"],
        ],
        [1600, 4300, 3100],
    )
    web_specs = [
        ("screen_01_home.png", "web_home_fit.jpg", "图 31  静态 Web 首页截图", "首屏给出样本范围、核心 KPI、主推荐和自动演示入口，便于答辩快速进入主题。"),
        ("screen_06_decision.png", "web_decision_fit.jpg", "图 32  决策页面截图", "决策页把 λ 权重、TOPSIS、风险准则和敏感性分析放在同一证据面板。"),
        ("screen_07_method.png", "web_method_fit.jpg", "图 33  方法说明与尾页截图", "方法页用完整公式块对应系统工程方法链，尾页给出条件式推荐和交付索引。"),
        ("screen_04_network.png", "web_network_fit.jpg", "图 34  网络页面截图", "网络页把关键机场、机场画像和传播关系集中展示，便于从图回到状态空间模型。"),
    ]
    web_figures = []
    for source_name, fitted_name, caption, reading in web_specs:
        fit_image_to_canvas(SCREEN / source_name, ASSETS / fitted_name, canvas=(1600, 900), padding=45)
        web_figures.append((fitted_name, caption, reading))
    add_picture_grid(
        doc,
        web_figures,
        source=ASSETS,
        width_in=2.55,
    )
    add_labeled_points(
        doc,
        "Web 截图与报告正文的对应关系",
        [
            ("首页", "首页不是单纯展示界面，而是把样本范围、关键 KPI 和策略推荐压缩成答辩开场面板，对应摘要中的核心结论。"),
            ("决策页", "决策页把 TOPSIS、风险准则和敏感性分析放在一起，便于现场说明为什么推荐不是绝对最优，而是有偏好条件的综合建议。"),
            ("方法页", "方法页对应第 2 章系统工程方法链，并把风险预测、状态空间、TOPSIS 和损失函数公式完整呈现，避免答辩时公式被截断或只剩符号片段。"),
            ("网络页", "网络页对应第 6-8 章，能从关键机场、机场画像和传播关系回到状态空间仿真输入。"),
            ("自动演示", "页面自动演示会真实改变控件状态：选择机场、点击航线气泡、调节风险滑杆、点击网络节点、切换情景策略和 λ 权重，保证讲稿、PPT 与网页动作一致。"),
        ],
    )
    add_table_from_rows(
        doc,
        "表 9  关键交付物索引",
        ["类别", "路径", "用途"],
        [
            ["Word/PDF 报告", "提交包/01_航空网络延误传播与恢复策略_报告*", "正式提交与审阅"],
            ["PPT", "提交包/02_航空网络延误传播与恢复策略_PPT.pptx", "课堂展示"],
            ["Web Demo", "提交包/04_FlightResilience_Static_Web/index.html", "交互式证据面板"],
            ["核心代码", "提交包/05_核心代码", "复现数据、模型、仿真和导出"],
            ["补充矩阵", "提交包/06_补充图表与矩阵", "ISM、AHP、传播矩阵与评价表"],
        ],
        [1500, 4300, 3200],
    )
    add_body(
        doc,
        """
        工程化交付的意义在于把报告结论变成可检查、可复现、可演示的证据系统。Word 报告承担正式论证，PPT 承担课堂展示，Web Demo 承担交互复核，补充图表与矩阵承担细节追溯，核心代码承担复现实验。这样即使读者不运行完整 Python 流程，也可以通过静态页面和补充矩阵检查每个结论的来源；如果需要复现，则可以从脚本和数据清单逐步回到原始处理链。
        """,
    )
    add_compact_matrix(
        doc,
        "工程化交付对应的反馈闭环",
        ["交付物", "反馈作用", "系统工程意义"],
        [
            ["Word 报告", "承载完整论证、方法映射、图文解释和结论边界。", "对应系统分析与系统评价的正式文档化。"],
            ["PPT/讲稿", "把复杂方法压缩为课堂可讲述的逻辑链。", "对应决策沟通与方案选择。"],
            ["静态 Web", "让图表、数据和方法页面可以现场交互查看。", "对应实施反馈和可视化复核。"],
            ["补充图表与矩阵", "保存 ISM、传播矩阵、权重、损失矩阵等细节。", "对应模型透明度和可追溯性。"],
            ["核心代码", "复现清洗、训练、网络、仿真、评价和导出。", "对应系统工程活动的可重复执行。"],
        ],
        [1700, 3600, 3700],
        body_size=7.0,
        header_size=7.3,
    )

    add_heading(doc, "11 结论、局限与后续改进", 1)
    add_heading(doc, "11.1 主要结论", 2)
    add_numbered(
        doc,
        [
            "航空延误在本项目样本中表现为时间、机场和航线共同作用的网络传播问题，不能只用单点平均值解释。",
            "计划阶段风险预测的主要价值在于提供传播和恢复策略的风险输入，而不是孤立替代调度决策。",
            "关键机场排序需要综合流量、中心性、预测风险和历史延误；MIA、DEN、DFW、FLL、MCO 是当前样本下的重点节点。",
            "状态空间模型在网络掩码和稳定化后可以用于情景推演，谱半径小于 1 保证设定情景下不会无限发散。",
            "动态组合在恢复优先偏好下表现最好，但风险型和不确定型决策显示基准策略在成本保守偏好下更稳妥。",
        ],
    )
    add_heading(doc, "11.2 局限与改进方向", 2)
    add_body(
        doc,
        """
        本项目仍有局限。第一，BTS 公开数据缺少完整实时天气、空域流控、机组和飞机轮转数据，因此传播路径只能用机场小时状态和航线连接近似。第二，恢复成本为相对情景参数，尚未接入真实运营成本。第三，预测模型 AUC 中等，说明计划阶段可知信息对极端延误仍有较大不确定性。第四，策略仿真采用有限情景，尚未覆盖更复杂的多扰动叠加。

        后续可在四个方向提升：引入 NOAA/METAR 天气与机场容量数据，增强外部扰动解释；引入飞机尾号和机组连接关系，刻画轮转传播；将恢复策略扩展为约束优化模型，显式处理预算、登机口、机组和航班优先级；将静态 Demo 升级为可选择情景参数的交互式仿真工具。
        """,
    )
    add_table_from_rows(
        doc,
        "局限、影响与改进路径",
        ["局限", "可能影响", "改进方向"],
        [
            ["缺少实时天气和容量数据", "外部扰动解释不足，weather 情景仍偏抽象。", "接入 METAR/NOAA 天气、机场容量、流控公告等数据。"],
            ["缺少飞机尾号和机组轮转", "传播路径只能用机场小时状态和航线连接近似。", "引入 tail number、机组连接和旅客中转关系，构建更细粒度传播网络。"],
            ["恢复成本为相对参数", "成本保守情景下的推荐边界只能作为策略比较。", "与真实运行成本、延误赔付、资源调配成本建立映射。"],
            ["预测模型指标中等", "计划阶段极端延误仍有较大不确定性。", "补充天气、容量、历史拥堵强度和更长时间窗口，并进行概率校准。"],
            ["仿真情景有限", "尚未覆盖多扰动叠加和滚动决策。", "将状态空间模型扩展为滚动优化或强化学习式策略仿真。"],
        ],
        [2500, 3200, 3300],
    )
    add_table_from_rows(
        doc,
        "表 10  多情景策略表现节选",
        ["情景", "策略", "累计延误", "恢复时间", "最低性能", "成本"],
        [
            [r.scenario, format_strategy(r.strategy), r.cumulative_delay, r.recovery_time, r.min_performance, r.strategy_cost]
            for r in scenarios.itertuples(index=False)
        ][:12],
        [1300, 1600, 1700, 1400, 1400, 1300],
    )
    add_labeled_points(
        doc,
        "结论的适用边界",
        [
            ("数据边界", "结论基于 2024Q1 BTS 公开数据和样本内前 30 个主要机场，适合做系统工程课程项目和策略比较，不等同于航空公司实时运行系统。"),
            ("模型边界", "预测模型使用计划阶段特征和训练期历史风险，状态空间模型使用机场小时状态和航线掩码，能够表达传播关系但不能覆盖所有机尾号、机组和天气细节。"),
            ("决策边界", "动态组合是恢复优先偏好下的综合推荐；当成本权重极高、概率不可知或决策者极端保守时，基准策略在损失矩阵准则下更稳妥。"),
            ("改进方向", "后续若接入实时天气、容量、尾号轮转和真实成本，可把当前离线策略比较升级为滚动优化或在线仿真决策。"),
        ],
    )

    add_heading(doc, "附录 A  复现路径与核心矩阵", 1)
    add_body(
        doc,
        """
        复现顺序建议为：`scripts/01_prepare_data.py` 完成数据准备；`02_train_model.py` 训练预测模型；`03_build_network.py` 构建机场网络；`04_fit_propagation.py` 估计传播矩阵；`05_run_simulation.py` 运行扰动仿真；`06_rank_strategies.py` 进行综合评价；`08_generate_report.py`、`09_generate_slides.py`、`11_export_static_web_assets.py` 生成提交物。核心矩阵包括 ISM 邻接矩阵、ISM 可达矩阵、传播矩阵 A、AHP 判断矩阵、损失矩阵和权重敏感性表。
        """,
    )
    add_table_from_rows(
        doc,
        "复现脚本与系统工程逻辑维对应",
        ["逻辑维步骤", "脚本/文件", "输出"],
        [
            ["摆明问题与资料收集", "01_prepare_data.py、data_manifest.csv", "清洗数据、审计表、训练/验证/测试切分。"],
            ["建立模型", "02_train_model.py、03_build_network.py、04_fit_propagation.py", "预测模型、机场网络、传播矩阵。"],
            ["方案综合与仿真", "05_run_simulation.py", "多情景策略结果、恢复曲线、传播范围。"],
            ["评价与决策", "06_rank_strategies.py", "TOPSIS、AHP（层次分析法）/熵权、模糊综合评价、风险与不确定决策。"],
            ["实施与反馈", "09_generate_slides.py、11_export_static_web_assets.py、12_upgrade_word_report.py", "PPT、Web Demo、升级版 Word 报告。"],
        ],
        [2200, 3600, 3200],
    )
    add_table_from_rows(
        doc,
        "表 11  模糊综合评价结果",
        ["策略", "优秀", "良好", "一般", "较差", "主导等级"],
        [[format_strategy(r.strategy), r.优秀, r.良好, r.一般, r.较差, r.dominant_grade] for r in fuzzy.itertuples(index=False)],
        [1500, 1100, 1100, 1100, 1100, 1600],
    )
    add_table_from_rows(
        doc,
        "资料来源与方法引用说明",
        ["类别", "来源", "在报告中的作用"],
        [
            ["课程方法", "第一章至第五章系统工程 PDF：概述、方法论、模型化、评价、决策。", "提供整体性原则、霍尔三维结构、切克兰德、解释结构模型化技术(ISM)、状态空间、AHP（层次分析法）、模糊综合评价、风险/不确定决策等方法依据。"],
            ["运行数据", "U.S. DOT BTS Airline On-Time Performance 2024 年 1-3 月公开数据。", "用于数据审计、风险预测、机场网络构建、延误状态聚合和多情景仿真。"],
            ["模型资产", "reports/tables、reports/figures、models、data/demo 中的矩阵、指标和图形。", "支撑表格、图片、仿真曲线、策略排名和 Web Demo 数据面板。"],
            ["交付资产", "提交包中的 Word/PDF、PPT、讲稿、静态 Web、核心代码与补充图表。", "用于课程提交、答辩展示、复现检查和现场追问时的证据回溯。"],
        ],
        [1500, 3500, 4000],
    )
    add_labeled_points(
        doc,
        "答辩时可强调的三句话",
        [
            ("第一句", "本项目不是单独做延误预测，而是把延误预测作为机场网络传播和恢复策略决策的风险输入。"),
            ("第二句", "系统工程方法体现在问题界定、系统边界、ISM 结构分析、状态空间动态模型、系统评价和风险/不确定决策的完整链路中。"),
            ("第三句", "最终推荐采用条件式表达：恢复效率和网络韧性优先时选动态组合，成本极端保守或概率不可知时保留基准策略。"),
        ],
    )
    add_table_from_rows(
        doc,
        "系统工程方法、图表证据与答辩口径速查",
        ["方法/概念", "报告位置", "图表证据", "一句话解释"],
        [
            ["霍尔三维结构", "第 2 章", "课程 PDF 使用清单、方法链说明", "用时间维、逻辑维和知识维组织从数据到决策的全过程。"],
            ["切克兰德软系统方法", "第 1-2 章", "多主体目标冲突、系统边界表", "把旅客、航司、机场和监管者的目标冲突显式写入评价体系。"],
            ["ISM 解释结构模型", "第 7 章", "图 1、图 17-18", "用邻接矩阵、可达矩阵和递阶结构解释延误传播因素。"],
            ["状态空间模型", "第 8 章", "图 19-24", "用 x(t+1)=A x(t)+B u(t)+G w(t)+epsilon(t) 进行扰动仿真。"],
            ["AHP/熵权/TOPSIS", "第 9 章", "表 5-6、图 25-28", "把恢复效果、韧性、影响、成本和复杂度统一成综合策略排序。"],
            ["风险/不确定决策", "第 9.1 章", "表 7-8、图 29-30", "检验不同概率假设和风险态度下推荐方案是否稳定。"],
        ],
        [1700, 1500, 1900, 4100],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    refresh_with_word(OUT)
    shutil.copy2(OUT, CANONICAL_OUT)
    return OUT


def main() -> None:
    out = build_report()
    print(out)


if __name__ == "__main__":
    main()
